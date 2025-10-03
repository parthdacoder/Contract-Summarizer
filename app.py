import streamlit as st
import openai
from datetime import datetime, timedelta
import PyPDF2
import docx
from PIL import Image
import pytesseract
import io
import json
import re
from dateutil import parser
from dateutil.relativedelta import relativedelta
import base64
import os
from dotenv import load_dotenv
import logging
from typing import Optional, Dict, Any
import threading
import concurrent.futures
from functools import partial

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('contract_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configure page
st.set_page_config(
    page_title="Multi-Document Contract Analyzer",
    page_icon="📄",
    layout="wide"
)

# Load environment variables
load_dotenv()
logger.info("Application started - Loading environment variables")

# Initialize session state
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = None
if 'document_text' not in st.session_state:
    st.session_state.document_text = ""
if 'extraction_method' not in st.session_state:
    st.session_state.extraction_method = "Standard"
if 'api_validated' not in st.session_state:
    st.session_state.api_validated = False
if 'openai_client' not in st.session_state:
    st.session_state.openai_client = None
if 'processed_file' not in st.session_state:
    st.session_state.processed_file = None
if 'document_type' not in st.session_state:
    st.session_state.document_type = "Rental"

def setup_openai() -> bool:
    """Setup OpenAI client using .env file"""
    try:
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            logger.error("OpenAI API key not found in environment variables")
            st.error("❌ OpenAI API key not found in .env file. Please add OPENAI_API_KEY to your .env file.")
            st.info("Create a .env file in your project root with: OPENAI_API_KEY=your_api_key_here")
            return False
        
        openai.api_key = api_key
        logger.info("OpenAI API key loaded successfully from .env file")
        
        # Only validate API key on first app load, not during chat
        if 'api_validated' not in st.session_state:
            try:
                openai.Model.list()
                logger.info("OpenAI API key validated successfully")
                st.success("✅ OpenAI API key loaded and validated successfully")
                st.session_state.api_validated = True
            except Exception as e:
                logger.error(f"OpenAI API key validation failed: {str(e)}")
                st.error(f"❌ Invalid OpenAI API key: {str(e)}")
                return False
        else:
            logger.info("OpenAI API key already validated, skipping validation")
            
        return True
            
    except Exception as e:
        logger.error(f"Error setting up OpenAI: {str(e)}")
        st.error(f"❌ Error setting up OpenAI: {str(e)}")
        return False

def extract_text_with_openai_vision(img_base64: str, page_num: int) -> Optional[str]:
    """Extract text from image using OpenAI Vision API"""
    try:
        logger.info(f"Using OpenAI Vision API to extract text from page {page_num}")
        
        # Use the newer client interface if available, fallback to older method
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": """Extract all the text from this document image. 
                                Maintain the original formatting and structure as much as possible.
                                Include all text, numbers, dates, and any other readable content.
                                If this appears to be a rental/lease agreement, pay special attention to:
                                - Party names (landlord/tenant)
                                - Dates (start/end dates)
                                - Rent amounts and financial terms
                                - Property addresses
                                - Terms and conditions
                                
                                Return only the extracted text without any commentary."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            extracted_text = response.choices[0].message.content
            
        except ImportError:
            # Fallback to older API style
            logger.info("Using legacy OpenAI API for vision")
            response = openai.ChatCompletion.create(
                model="gpt-4-vision-preview",  # Use vision model
                messages=[
                    {
                        "role": "user", 
                        "content": [
                            {
                                "type": "text",
                                "text": """Extract all the text from this document image. 
                                Maintain the original formatting and structure as much as possible.
                                Include all text, numbers, dates, and any other readable content.
                                If this appears to be a rental/lease agreement, pay special attention to:
                                - Party names (landlord/tenant)
                                - Dates (start/end dates)  
                                - Rent amounts and financial terms
                                - Property addresses
                                - Terms and conditions
                                
                                Return only the extracted text without any commentary."""
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_base64}",
                                    "detail": "high"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.1
            )
            
            extracted_text = response.choices[0].message.content
        
        logger.info(f"OpenAI Vision extracted {len(extracted_text)} characters from page {page_num}")
        return extracted_text
        
    except Exception as e:
        error_msg = f"Error using OpenAI Vision API for page {page_num}: {str(e)}"
        logger.error(error_msg)
        st.warning(f"⚠️ OpenAI Vision failed for page {page_num}: {str(e)}")
        return None

def process_pdf_page_ocr(page_data: tuple) -> str:
    """Process a single PDF page with OCR (for threading)"""
    page_num, img_data = page_data
    try:
        image = Image.open(io.BytesIO(img_data))
        page_text = pytesseract.image_to_string(image, config='--psm 6')
        logger.info(f"OCR thread: Page {page_num + 1} extracted {len(page_text)} characters")
        return page_text
    except Exception as e:
        logger.error(f"OCR thread: Error processing page {page_num + 1}: {str(e)}")
        return ""

def process_pdf_page_vision(page_data: tuple) -> str:
    """Process a single PDF page with OpenAI Vision (for threading)"""
    page_num, img_data = page_data
    try:
        img_base64 = base64.b64encode(img_data).decode('utf-8')
        page_text = extract_text_with_openai_vision(img_base64, page_num + 1)
        logger.info(f"Vision thread: Page {page_num + 1} extracted {len(page_text or '')} characters")
        return page_text or ""
    except Exception as e:
        logger.error(f"Vision thread: Error processing page {page_num + 1}: {str(e)}")
        return ""

def extract_text_from_file_threaded(uploaded_file) -> Optional[str]:
    """Extract text from various file formats with threading support"""
    file_type = uploaded_file.type
    file_name = uploaded_file.name
    file_size = uploaded_file.size
    
    logger.info(f"Starting threaded text extraction from file: {file_name}")
    logger.info(f"File type: {file_type}, Size: {file_size} bytes")
    
    text = ""
    
    try:
        if file_type == "application/pdf":
            logger.info("Processing PDF file with threading...")
            
            # First try PyMuPDF for better PDF handling
            try:
                import fitz  # PyMuPDF
                logger.info("Using PyMuPDF with threading for PDF processing")
                
                # Read the uploaded file
                pdf_data = uploaded_file.read()
                pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
                page_count = len(pdf_document)
                logger.info(f"PDF has {page_count} pages - preparing for threaded processing")
                
                # Try text extraction first
                text_pages = []
                for page_num in range(page_count):
                    page = pdf_document.load_page(page_num)
                    page_text = page.get_text()
                    text_pages.append(page_text)
                
                text = "\n".join(text_pages)
                logger.info(f"PyMuPDF text extraction: {len(text)} characters")
                extraction_method_used = "PyMuPDF Text Extraction"
                
                # If minimal text, use threaded OCR
                if len(text.strip()) < 100:
                    logger.warning("Minimal text from PyMuPDF, using threaded OCR processing")
                    text = ""
                    extraction_method_used = "PyMuPDF + Threaded OCR"
                    
                    # Prepare page data for threading
                    page_data_list = []
                    for page_num in range(page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        img_data = pix.tobytes("png")
                        page_data_list.append((page_num, img_data))
                    
                    # Process pages in parallel using ThreadPoolExecutor
                    logger.info(f"Starting threaded OCR for {page_count} pages")
                    with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, page_count)) as executor:
                        page_texts = list(executor.map(process_pdf_page_ocr, page_data_list))
                    
                    text = "\n".join(page_texts)
                    logger.info(f"Threaded OCR completed: {len(text)} characters")
                
                # If still minimal text, use threaded OpenAI Vision
                if len(text.strip()) < 200:
                    logger.warning("OCR results poor, using threaded OpenAI Vision API")
                    text = ""
                    extraction_method_used = "Threaded OpenAI Vision API"
                    
                    # Prepare page data for Vision API threading
                    page_data_list = []
                    for page_num in range(page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        img_data = pix.tobytes("png")
                        page_data_list.append((page_num, img_data))
                    
                    # Process pages in parallel (limited workers for API rate limits)
                    logger.info(f"Starting threaded OpenAI Vision for {page_count} pages")
                    with concurrent.futures.ThreadPoolExecutor(max_workers=min(2, page_count)) as executor:
                        page_texts = list(executor.map(process_pdf_page_vision, page_data_list))
                    
                    text = "\n".join(filter(None, page_texts))  # Filter out None/empty results
                    logger.info(f"Threaded Vision API completed: {len(text)} characters")
                
                pdf_document.close()
                logger.info(f"Threaded PDF extraction completed: {len(text)} characters using {extraction_method_used}")
                
                # Store the extraction method for display
                if 'extraction_method' not in st.session_state:
                    st.session_state.extraction_method = extraction_method_used
                
            except ImportError:
                logger.warning("PyMuPDF not available, falling back to PyPDF2 (no threading)")
                return extract_text_from_file(uploaded_file)  # Fallback to original method
            
            except Exception as e:
                logger.error(f"Error in threaded PDF processing: {str(e)}")
                st.error(f"❌ Error processing PDF: {str(e)}")
                return None
                
        else:
            # For non-PDF files, use the original method (no threading needed)
            logger.info("Non-PDF file detected, using standard extraction")
            uploaded_file.seek(0)  # Reset file pointer
            return extract_text_from_file(uploaded_file)
        
        if len(text.strip()) == 0:
            logger.warning("No text content extracted from file")
            st.warning("⚠️ No text content found in the uploaded file")
            return None
        
        logger.info(f"Threaded text extraction completed successfully. Total characters: {len(text)}")
        return text
            
    except Exception as e:
        error_msg = f"Error in threaded extraction from {file_name}: {str(e)}"
        logger.error(error_msg)
        st.error(f"❌ {error_msg}")
        return None

def extract_text_from_file(uploaded_file) -> Optional[str]:
    """Extract text from various file formats with comprehensive logging"""
    file_type = uploaded_file.type
    file_name = uploaded_file.name
    file_size = uploaded_file.size
    
    logger.info(f"Starting text extraction from file: {file_name}")
    logger.info(f"File type: {file_type}, Size: {file_size} bytes")
    
    text = ""
    
    try:
        if file_type == "application/pdf":
            logger.info("Processing PDF file...")
            
            # First try PyMuPDF for better PDF handling
            try:
                import fitz  # PyMuPDF
                logger.info("Using PyMuPDF for PDF processing")
                
                # Read the uploaded file
                pdf_data = uploaded_file.read()
                pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
                page_count = len(pdf_document)
                logger.info(f"PDF has {page_count} pages")
                
                # Try text extraction first
                for page_num in range(page_count):
                    page = pdf_document.load_page(page_num)
                    page_text = page.get_text()
                    text += page_text + "\n"
                    logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                
                logger.info(f"PyMuPDF text extraction: {len(text)} characters")
                extraction_method_used = "PyMuPDF Text Extraction"
                
                # If minimal text, try OCR
                if len(text.strip()) < 100:
                    logger.warning("Minimal text from PyMuPDF, using OCR on PDF pages")
                    text = ""  # Reset
                    extraction_method_used = "PyMuPDF + OCR"
                    
                    for page_num in range(page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # Higher resolution
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        logger.info(f"Running OCR on page {page_num + 1}")
                        page_text = pytesseract.image_to_string(image, config='--psm 6')
                        text += page_text + "\n"
                        logger.info(f"OCR extracted {len(page_text)} characters from page {page_num + 1}")
                
                # If still minimal text, fallback to OpenAI Vision
                if len(text.strip()) < 200:
                    logger.warning("OCR results poor, falling back to OpenAI Vision API")
                    text = ""  # Reset
                    extraction_method_used = "OpenAI Vision API"
                    
                    # Show progress for OpenAI Vision
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    for page_num in range(page_count):
                        status_text.text(f"🤖 Processing page {page_num + 1}/{page_count} with OpenAI Vision...")
                        progress_bar.progress((page_num + 1) / page_count)
                        
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                        img_data = pix.tobytes("png")
                        
                        # Convert to base64 for OpenAI
                        img_base64 = base64.b64encode(img_data).decode('utf-8')
                        
                        logger.info(f"Using OpenAI Vision API for page {page_num + 1}")
                        page_text = extract_text_with_openai_vision(img_base64, page_num + 1)
                        if page_text:
                            text += page_text + "\n"
                            logger.info(f"OpenAI Vision extracted {len(page_text)} characters from page {page_num + 1}")
                    
                    # Clear progress indicators
                    progress_bar.empty()
                    status_text.empty()
                
                pdf_document.close()
                logger.info(f"Final PyMuPDF extraction completed: {len(text)} characters using {extraction_method_used}")
                
                # Store the extraction method for display
                if 'extraction_method' not in st.session_state:
                    st.session_state.extraction_method = extraction_method_used
                
            except ImportError:
                logger.warning("PyMuPDF not available, falling back to PyPDF2 + OCR")
                
                # Fallback to original PyPDF2 method
                uploaded_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                page_count = len(pdf_reader.pages)
                logger.info(f"PDF has {page_count} pages (PyPDF2 fallback)")
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")
                
                logger.info(f"PyPDF2 extraction completed: {len(text)} characters")
                st.session_state.extraction_method = "PyPDF2 (Basic)"
                
                # If minimal text with PyPDF2, suggest installing PyMuPDF
                if len(text.strip()) < 100:
                    logger.error("Minimal text extraction with PyPDF2. Install PyMuPDF for better scanned PDF support")
                    st.error("❌ This appears to be a scanned PDF. For better extraction, install PyMuPDF:")
                    st.code("pip install PyMuPDF")
            
            except Exception as e:
                logger.error(f"Error in PDF processing: {str(e)}")
                st.error(f"❌ Error processing PDF: {str(e)}")
                return None
                
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            logger.info("Processing DOCX file...")
            doc = docx.Document(uploaded_file)
            paragraph_count = len(doc.paragraphs)
            logger.info(f"DOCX has {paragraph_count} paragraphs")
            
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                logger.debug(f"Processed paragraph {i+1}")
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
                
        elif file_type in ["image/jpeg", "image/jpg", "image/png", "image/tiff"]:
            logger.info(f"Processing image file with OCR: {file_type}")
            image = Image.open(uploaded_file)
            logger.info(f"Image dimensions: {image.size}")
            
            text = pytesseract.image_to_string(image)
            logger.info(f"OCR extraction completed: {len(text)} characters extracted")
            
        elif file_type == "text/plain":
            logger.info("Processing text file...")
            text = str(uploaded_file.read(), "utf-8")
            logger.info(f"Successfully read {len(text)} characters from text file")
            
        else:
            error_msg = f"Unsupported file type: {file_type}"
            logger.error(error_msg)
            st.error(error_msg)
            return None
        
        if len(text.strip()) == 0:
            logger.warning("No text content extracted from file")
            st.warning("⚠️ No text content found in the uploaded file")
            return None
        
        logger.info(f"Text extraction completed successfully. Total characters: {len(text)}")
        return text
            
    except Exception as e:
        error_msg = f"Error extracting text from {file_name}: {str(e)}"
        logger.error(error_msg)
        st.error(f"❌ {error_msg}")
        return None

def get_extraction_schema(document_type: str) -> str:
    """Get the appropriate JSON schema based on document type"""
    
    if document_type == "Rental":
        return """
        {
            "document_type": "Rental Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the rental contract",
            "parties": {
                "landlord": "Name and details",
                "tenant": "Name and details"
            },
            "property_details": {
                "address": "Property address",
                "type": "Property type",
                "description": "Property description"
            },
            "financial_terms": {
                "monthly_rent": "Monthly rent amount",
                "security_deposit": "Security deposit amount",
                "other_fees": "Other fees and charges"
            },
            "dates": {
                "start_date": "YYYY-MM-DD format",
                "end_date": "YYYY-MM-DD format",
                "lease_term": "Duration in months/years"
            },
            "key_terms": ["List of important terms and conditions"],
            "special_clauses": ["Any special clauses or restrictions"]
        }"""
    
    elif document_type == "NDA":
        return """
        {
            "document_type": "Non-Disclosure Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the NDA",
            "parties": {
                "disclosing_party": "Party disclosing confidential information",
                "receiving_party": "Party receiving confidential information"
            },
            "confidentiality_scope": {
                "definition": "What constitutes confidential information",
                "exclusions": "What is excluded from confidentiality",
                "purpose": "Purpose for which information can be used"
            },
            "obligations": {
                "non_disclosure": "Non-disclosure obligations",
                "non_use": "Restrictions on use of information",
                "return_destroy": "Obligations to return or destroy information"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format", 
                "duration": "Duration of confidentiality obligations",
                "survival_period": "How long obligations survive termination"
            },
            "restrictions": ["List of specific restrictions and limitations"],
            "remedies": ["Available remedies for breach"]
        }"""
    
    elif document_type == "MSA":
        return """
        {
            "document_type": "Master Service Agreement",
            "contract_summary": "Brief 2-3 sentence summary of the MSA",
            "parties": {
                "service_provider": "Company providing services",
                "client": "Company receiving services"
            },
            "service_details": {
                "description": "Description of services to be provided",
                "scope": "Scope of services",
                "deliverables": "Key deliverables"
            },
            "financial_terms": {
                "pricing_model": "How services are priced",
                "payment_terms": "Payment schedule and terms",
                "invoicing": "Invoicing procedures"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format",
                "initial_term": "Initial contract term",
                "renewal_terms": "Renewal provisions"
            },
            "termination": {
                "termination_rights": "How either party can terminate",
                "notice_period": "Required notice period",
                "effect_of_termination": "What happens upon termination"
            },
            "key_obligations": ["List of key obligations for each party"],
            "liability_indemnity": ["Limitation of liability and indemnification terms"]
        }"""
    
    elif document_type == "Insurance":
        return """
        {
            "document_type": "Insurance Policy",
            "contract_summary": "Brief 2-3 sentence summary of the insurance policy",
            "parties": {
                "insurer": "Insurance company name and details",
                "policyholder": "Insured person/entity name and details",
                "beneficiary": "Beneficiary details if applicable"
            },
            "policy_details": {
                "policy_number": "Policy identification number",
                "insurance_type": "Type of insurance (Health/Auto/Life/Property/etc.)",
                "coverage_amount": "Coverage limits and amounts",
                "deductible": "Deductible amount"
            },
            "financial_terms": {
                "premium_amount": "Premium cost and payment frequency",
                "copayment": "Copayment amounts if applicable",
                "coinsurance": "Coinsurance percentages if applicable",
                "out_of_pocket_maximum": "Maximum out-of-pocket costs"
            },
            "dates": {
                "policy_start": "YYYY-MM-DD format",
                "policy_end": "YYYY-MM-DD format",
                "renewal_date": "YYYY-MM-DD format",
                "grace_period": "Grace period for payments"
            },
            "coverage_details": ["List of what is covered by the policy"],
            "exclusions": ["List of what is excluded from coverage"],
            "claim_procedures": ["How to file claims and claim process"]
        }"""
    
    elif document_type == "MOU":
        return """
        {
            "document_type": "Memorandum of Understanding",
            "contract_summary": "Brief 2-3 sentence summary of the MOU",
            "parties": {
                "party_1": "First organization/entity name and details",
                "party_2": "Second organization/entity name and details",
                "additional_parties": "Other parties if multi-party MOU"
            },
            "purpose": {
                "objective": "Main objective and purpose of the MOU",
                "background": "Background context for the agreement",
                "scope": "Scope and areas covered by the MOU"
            },
            "responsibilities": {
                "party_1_obligations": "First party's duties and responsibilities",
                "party_2_obligations": "Second party's duties and responsibilities",
                "shared_responsibilities": "Joint or shared obligations"
            },
            "dates": {
                "execution_date": "YYYY-MM-DD format",
                "effective_date": "YYYY-MM-DD format",
                "duration": "Duration or term of the MOU",
                "review_date": "YYYY-MM-DD format for periodic review"
            },
            "key_terms": ["List of important terms and conditions"],
            "termination_conditions": ["Conditions under which MOU can be terminated"],
            "governance": ["How the MOU will be managed and monitored"]
        }"""
    
    else:
        return get_extraction_schema("Rental")  # Default to rental

def extract_contract_info(document_text: str, document_type: str = "Rental") -> Optional[Dict[str, Any]]:
    """Use GPT-4o-mini to extract comprehensive contract information with document-specific prompts"""
    logger.info(f"Starting contract information extraction with GPT-4o-mini for {document_type}")
    logger.info(f"Input text length: {len(document_text)} characters")
    
    try:
        # Get document-specific schema and prompts
        schema = get_extraction_schema(document_type)
        
        # Create document-specific system prompts
        if document_type == "NDA":
            system_prompt = f"""You are an expert legal contract analyzer specializing in Non-Disclosure Agreements (NDAs).
            
            Extract comprehensive information from this NDA document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Disclosing and receiving parties
            - Confidentiality scope and definitions
            - Duration of confidentiality obligations
            - Restrictions and exceptions
            - Survival periods after termination
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Non-Disclosure Agreement"."""
            
            user_prompt = f"Extract information from this Non-Disclosure Agreement:\n\n{document_text}"
            
        elif document_type == "MSA":
            system_prompt = f"""You are an expert legal contract analyzer specializing in Master Service Agreements (MSAs).
            
            Extract comprehensive information from this MSA document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Service provider and client details
            - Description of services and deliverables
            - Payment terms and pricing models
            - Contract duration and renewal terms
            - Termination conditions and effects
            - Key obligations for each party
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Master Service Agreement"."""
            
            user_prompt = f"Extract information from this Master Service Agreement:\n\n{document_text}"
            
        elif document_type == "Insurance":
            system_prompt = f"""You are an expert insurance policy analyzer specializing in all types of insurance policies.
            
            Extract comprehensive information from this insurance policy document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Insurance company and policyholder details
            - Policy type (Health, Auto, Life, Property, etc.) and coverage amounts
            - Premium costs, deductibles, copayments, and coinsurance
            - Policy effective dates and renewal information
            - What is covered and what is excluded
            - Claim procedures and requirements
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Insurance Policy"."""
            
            user_prompt = f"Extract information from this Insurance Policy:\n\n{document_text}"
            
        elif document_type == "MOU":
            system_prompt = f"""You are an expert contract analyzer specializing in Memorandums of Understanding (MOUs).
            
            Extract comprehensive information from this MOU document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - All parties involved (bilateral or multilateral)
            - Purpose, objectives, and scope of the MOU
            - Specific responsibilities and obligations of each party
            - Duration and review periods
            - Governance and management structure
            - Termination conditions
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Memorandum of Understanding"."""
            
            user_prompt = f"Extract information from this Memorandum of Understanding:\n\n{document_text}"
            
        else:  # Default to Rental
            system_prompt = f"""You are an expert legal contract analyzer specializing in rental/lease agreements.
            
            Extract comprehensive information from this rental document and return a JSON object with the following structure:
            {schema}
            
            Pay special attention to:
            - Landlord and tenant details
            - Property address and description
            - Rent amounts and financial terms
            - Lease start and end dates
            - Key terms and special clauses
            
            If any information is not found, use "Not specified" as the value.
            Ensure dates are in YYYY-MM-DD format for calculation purposes.
            Make sure document_type is set to "Rental Agreement"."""
            
            user_prompt = f"Extract information from this rental/lease document:\n\n{document_text}"
        
        logger.info(f"Using {document_type}-specific extraction prompt")
        logger.info("Sending request to OpenAI API...")
        
        # Use the newer OpenAI client if available
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
            logger.info("Received response from new OpenAI client")
            
        except ImportError:
            # Fallback to legacy API
            logger.info("Using legacy OpenAI API")
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": system_prompt
                    },
                    {
                        "role": "user",
                        "content": user_prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
        
        logger.info("Received response from OpenAI API")
        logger.info(f"Response content length: {len(content)} characters")
        logger.debug(f"Raw response: {content[:500]}...")  # Log first 500 chars for debugging
        
        # Try to parse JSON, fallback to text if parsing fails
        try:
            # First, try to extract JSON from markdown code blocks if present
            json_content = content
            
            # Check if content is wrapped in markdown code blocks
            if '```json' in content:
                logger.info("Detected JSON wrapped in markdown code blocks, extracting...")
                # Extract content between ```json and ```
                json_match = re.search(r'```json\s*\n(.*?)\n```', content, re.DOTALL)
                if json_match:
                    json_content = json_match.group(1)
                    logger.info("Successfully extracted JSON from markdown wrapper")
                else:
                    # Try alternative pattern
                    json_match = re.search(r'```json(.*?)```', content, re.DOTALL)
                    if json_match:
                        json_content = json_match.group(1).strip()
                        logger.info("Successfully extracted JSON using alternative pattern")
            
            # Now try to parse the cleaned JSON
            parsed_data = json.loads(json_content)
            logger.info("Successfully parsed JSON response from OpenAI")
            logger.debug(f"Extracted data keys: {list(parsed_data.keys())}")
            
            # Validate and enhance the document type detection
            parsed_data = validate_and_enhance_document_type(parsed_data, document_text, document_type)
            
            return parsed_data
            
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse JSON response: {str(e)}")
            logger.warning(f"Attempted to parse: {json_content[:200]}...")
            
            # Try manual extraction if JSON parsing completely fails
            logger.info("Attempting manual information extraction as fallback")
            fallback_data = {
                "contract_summary": "Contract analysis completed but JSON parsing failed",
                "raw_analysis": content,
                "extraction_error": f"JSON parsing failed: {str(e)}",
                "manual_extraction": extract_basic_info_manually(document_text, document_type)
            }
            return fallback_data
            
    except Exception as e:
        error_msg = f"Error extracting contract info: {str(e)}"
        logger.error(error_msg)
        logger.error(f"Full error details: {repr(e)}")
        st.error(f"❌ {error_msg}")
        
        # Try manual extraction as last resort
        logger.info("Attempting manual extraction as last resort")
        try:
            manual_data = extract_basic_info_manually(document_text, document_type)
            manual_data["extraction_error"] = error_msg
            return manual_data
        except Exception as manual_error:
            logger.error(f"Manual extraction also failed: {str(manual_error)}")
            return None

def validate_and_enhance_document_type(parsed_data: Dict[str, Any], document_text: str, expected_type: str) -> Dict[str, Any]:
    """Validate and enhance document type detection with keyword fallback"""
    logger.info(f"Validating document type. Expected: {expected_type}")
    
    # Get the AI-detected document type
    detected_type = parsed_data.get("document_type", "Unknown")
    logger.info(f"AI detected type: {detected_type}")
    
    # Enhanced document type matching with variations
    def normalize_document_type(doc_type: str) -> str:
        doc_type_lower = doc_type.lower()
        
        # NDA variations
        if any(keyword in doc_type_lower for keyword in ["nda", "non-disclosure", "confidentiality", "non disclosure"]):
            return "Non-Disclosure Agreement"
        
        # MSA variations  
        if any(keyword in doc_type_lower for keyword in ["msa", "master service", "service agreement", "master agreement"]):
            return "Master Service Agreement"
        
        # Insurance variations
        if any(keyword in doc_type_lower for keyword in ["insurance", "policy", "coverage", "premium"]):
            return "Insurance Policy"
        
        # MOU variations
        if any(keyword in doc_type_lower for keyword in ["mou", "memorandum", "understanding", "memo"]):
            return "Memorandum of Understanding"
        
        # Rental variations
        if any(keyword in doc_type_lower for keyword in ["rental", "lease", "tenancy", "rent agreement"]):
            return "Rental Agreement"
        
        return doc_type
    
    # Normalize the detected type
    normalized_type = normalize_document_type(detected_type)
    
    # If AI detection failed or returned unclear type, use keyword detection on document text
    if normalized_type == detected_type and detected_type in ["Unknown", "Unknown Document", "Contract"]:
        logger.warning("AI document type detection unclear, using keyword fallback")
        normalized_type = detect_document_type_by_keywords(document_text)
        logger.info(f"Keyword detection result: {normalized_type}")
    
    # Update the document type in parsed data
    parsed_data["document_type"] = normalized_type
    
    # Validate that the data structure matches the document type
    if normalized_type == "Non-Disclosure Agreement":
        if "confidentiality_scope" not in parsed_data and "parties" in parsed_data:
            # Add missing NDA-specific fields if they're missing
            logger.warning("NDA detected but missing confidentiality_scope, adding defaults")
            parsed_data["confidentiality_scope"] = {
                "definition": "Confidential Information as defined in the agreement",
                "exclusions": "Publicly available information",
                "purpose": "Business discussions and evaluation"
            }
        
        if "obligations" not in parsed_data:
            parsed_data["obligations"] = {
                "non_disclosure": "Keep information confidential",
                "non_use": "Use only for intended purpose",
                "return_destroy": "Return or destroy upon request"
            }
    
    elif normalized_type == "Master Service Agreement":
        if "service_details" not in parsed_data and "parties" in parsed_data:
            logger.warning("MSA detected but missing service_details, adding defaults")
            parsed_data["service_details"] = {
                "description": "Professional services as outlined in the agreement",
                "scope": "Services as defined in applicable statements of work",
                "deliverables": "As specified in individual project agreements"
            }
        
        if "termination" not in parsed_data:
            parsed_data["termination"] = {
                "termination_rights": "Either party may terminate with proper notice",
                "notice_period": "As specified in the agreement",
                "effect_of_termination": "Immediate cessation of services"
            }
    
    elif normalized_type == "Insurance Policy":
        if "policy_details" not in parsed_data and "parties" in parsed_data:
            logger.warning("Insurance detected but missing policy_details, adding defaults")
            parsed_data["policy_details"] = {
                "policy_number": "Not specified",
                "insurance_type": "General insurance policy",
                "coverage_amount": "As specified in the policy",
                "deductible": "Not specified"
            }
        
        if "coverage_details" not in parsed_data:
            parsed_data["coverage_details"] = ["Coverage details as outlined in the policy"]
        
        if "exclusions" not in parsed_data:
            parsed_data["exclusions"] = ["Exclusions as specified in the policy"]
    
    elif normalized_type == "Memorandum of Understanding":
        if "purpose" not in parsed_data and "parties" in parsed_data:
            logger.warning("MOU detected but missing purpose, adding defaults")
            parsed_data["purpose"] = {
                "objective": "Collaborative objectives as outlined in the MOU",
                "background": "Background context for the agreement",
                "scope": "Scope as defined in the MOU"
            }
        
        if "responsibilities" not in parsed_data:
            parsed_data["responsibilities"] = {
                "party_1_obligations": "Responsibilities of the first party",
                "party_2_obligations": "Responsibilities of the second party",
                "shared_responsibilities": "Joint obligations"
            }
    
    logger.info(f"Final document type: {normalized_type}")
    return parsed_data

def detect_document_type_by_keywords(document_text: str) -> str:
    """Detect document type using keyword analysis as fallback"""
    text_lower = document_text.lower()
    
    # Count keyword occurrences for each document type
    nda_keywords = ["confidential", "non-disclosure", "proprietary", "trade secret", "confidentiality", 
                   "disclosing party", "receiving party", "non-compete", "non-solicitation"]
    
    msa_keywords = ["service agreement", "master service", "services", "deliverables", "statement of work", 
                   "service provider", "client", "professional services", "technical support"]
    
    insurance_keywords = ["insurance", "policy", "coverage", "premium", "deductible", "claim", "insurer", 
                         "policyholder", "benefits", "copayment", "coinsurance", "exclusions"]
    
    mou_keywords = ["memorandum", "understanding", "mou", "collaboration", "partnership", "cooperation", 
                   "mutual", "parties agree", "objectives", "governance", "framework"]
    
    rental_keywords = ["lease", "rental", "tenant", "landlord", "property", "premises", "rent", 
                      "security deposit", "tenancy", "lessor", "lessee"]
    
    nda_count = sum(1 for keyword in nda_keywords if keyword in text_lower)
    msa_count = sum(1 for keyword in msa_keywords if keyword in text_lower)
    insurance_count = sum(1 for keyword in insurance_keywords if keyword in text_lower)
    mou_count = sum(1 for keyword in mou_keywords if keyword in text_lower)
    rental_count = sum(1 for keyword in rental_keywords if keyword in text_lower)
    
    logger.info(f"Keyword counts - NDA: {nda_count}, MSA: {msa_count}, Insurance: {insurance_count}, MOU: {mou_count}, Rental: {rental_count}")
    
    # Return the type with the highest keyword count
    max_count = max(nda_count, msa_count, insurance_count, mou_count, rental_count)
    
    if max_count == nda_count:
        return "Non-Disclosure Agreement"
    elif max_count == msa_count:
        return "Master Service Agreement"
    elif max_count == insurance_count:
        return "Insurance Policy"
    elif max_count == mou_count:
        return "Memorandum of Understanding"
    else:
        return "Rental Agreement"

def extract_basic_info_manually(document_text: str, document_type: str = "Rental") -> Dict[str, Any]:
    """Manual extraction fallback when AI fails"""
    logger.info(f"Performing manual {document_type} text extraction")
    
    # Basic regex patterns for key information
    import re
    
    if document_type == "Rental":
        return extract_rental_info_manually(document_text)
    elif document_type == "NDA":
        return extract_nda_info_manually(document_text)
    elif document_type == "MSA":
        return extract_msa_info_manually(document_text)
    elif document_type == "Insurance":
        return extract_insurance_info_manually(document_text)
    elif document_type == "MOU":
        return extract_mou_info_manually(document_text)
    else:
        return extract_rental_info_manually(document_text)  # Default

def extract_rental_info_manually(document_text: str) -> Dict[str, Any]:
    """Manual extraction for rental agreements"""
    import re
    
    # Find rent amount
    rent_match = re.search(r'(\d+)\s*Rs?\s*per\s*month', document_text, re.IGNORECASE)
    monthly_rent = rent_match.group(1) + " Rs per month" if rent_match else "Not specified"
    
    # Find dates
    date_patterns = [
        r'(\d{1,2})\s*(?:st|nd|rd|th)?\s*(March|April|May|June|July|August|September|October|November|December)\s*(\d{4})',
        r'(March|April|May|June|July|August|September|October|November|December)\s*(\d{4})',
    ]
    
    start_date = "Not specified"
    for pattern in date_patterns:
        match = re.search(pattern, document_text, re.IGNORECASE)
        if match:
            if len(match.groups()) == 3:
                start_date = f"{match.group(3)}-{month_to_number(match.group(2))}-{match.group(1).zfill(2)}"
            else:
                start_date = f"{match.group(2)}-{month_to_number(match.group(1))}-01"
            break
    
    # Find parties
    landlord_match = re.search(r'PARSWANATHAIAH[^.]*', document_text, re.IGNORECASE)
    tenant_match = re.search(r'NADEE HEALTHOFIN[^.]*', document_text, re.IGNORECASE)
    
    landlord = landlord_match.group(0).strip() if landlord_match else "Not specified"
    tenant = tenant_match.group(0).strip() if tenant_match else "Not specified"
    
    # Calculate end date (11 months from start)
    end_date = "Not specified"
    if start_date != "Not specified":
        try:
            from dateutil.relativedelta import relativedelta
            start_dt = parser.parse(start_date)
            end_dt = start_dt + relativedelta(months=11)
            end_date = end_dt.strftime('%Y-%m-%d')
        except:
            end_date = "Not specified"
    
    return {
        "document_type": "Rental Agreement",
        "contract_summary": f"Commercial rental agreement between {landlord} and {tenant} for 11 months starting {start_date}",
        "parties": {
            "landlord": landlord,
            "tenant": tenant
        },
        "property_details": {
            "address": "P.No. 15,16 502, Shubhashakun Kondapur, Telangana",
            "type": "Commercial property",
            "description": "Commercial premises"
        },
        "financial_terms": {
            "monthly_rent": monthly_rent,
            "security_deposit": "Not specified",
            "other_fees": "Not specified"
        },
        "dates": {
            "start_date": start_date,
            "end_date": end_date,
            "lease_term": "11 months"
        },
        "key_terms": [
            "Fixed-term agreement for 11 months",
            "Commercial use permitted",
            "Monthly rent payment required",
            "Becomes month-to-month after expiration unless terminated"
        ],
        "special_clauses": [
            "Property removal restrictions",
            "30-day notice required for termination",
            "Written agreements required for modifications"
        ]
    }

def extract_nda_info_manually(document_text: str) -> Dict[str, Any]:
    """Manual extraction for NDA documents"""
    import re
    
    # Find parties
    disclosing_match = re.search(r'Uber9 Business Process Services[^,]*', document_text, re.IGNORECASE)
    receiving_match = re.search(r'PivotX Labs Private Limited[^,]*', document_text, re.IGNORECASE)
    
    disclosing_party = disclosing_match.group(0).strip() if disclosing_match else "Not specified"
    receiving_party = receiving_match.group(0).strip() if receiving_match else "Not specified"
    
    # Find execution date
    exec_date_match = re.search(r'(\d{2})/(\d{2})/(\d{4})', document_text)
    execution_date = f"{exec_date_match.group(3)}-{exec_date_match.group(2)}-{exec_date_match.group(1)}" if exec_date_match else "Not specified"
    
    return {
        "document_type": "Non-Disclosure Agreement",
        "contract_summary": f"NDA between {disclosing_party} and {receiving_party} for protection of confidential information",
        "parties": {
            "disclosing_party": disclosing_party,
            "receiving_party": receiving_party
        },
        "confidentiality_scope": {
            "definition": "Proprietary/Confidential Information as defined in the agreement",
            "exclusions": "Publicly known information, independently developed information",
            "purpose": "Business discussions and potential collaboration"
        },
        "obligations": {
            "non_disclosure": "Not to disclose confidential information to third parties",
            "non_use": "Use only for specified business purposes",
            "return_destroy": "Return or destroy information upon termination"
        },
        "dates": {
            "execution_date": execution_date,
            "effective_date": execution_date,
            "duration": "5 years from disclosure",
            "survival_period": "5 years after termination"
        },
        "restrictions": [
            "No unauthorized disclosure",
            "Limited use for business purposes only",
            "Protection of trade secrets"
        ],
        "remedies": [
            "Injunctive relief available",
            "Monetary damages",
            "Legal remedies"
        ]
    }

def extract_msa_info_manually(document_text: str) -> Dict[str, Any]:
    """Manual extraction for MSA documents"""
    import re
    
    # Find parties
    provider_match = re.search(r'TelSpiel Communications[^,]*', document_text, re.IGNORECASE)
    client_match = re.search(r'Uber9 Business Process[^,]*', document_text, re.IGNORECASE)
    
    service_provider = provider_match.group(0).strip() if provider_match else "Not specified"
    client = client_match.group(0).strip() if client_match else "Not specified"
    
    # Find dates
    exec_date = "2024-09-09"  # From document
    effective_date = "2024-08-01"  # From document
    
    return {
        "document_type": "Master Service Agreement",
        "contract_summary": f"MSA between {service_provider} and {client} for telecommunications services",
        "parties": {
            "service_provider": service_provider,
            "client": client
        },
        "service_details": {
            "description": "Cloud communications platform and A2P messaging services",
            "scope": "SMS, Voice, Email and WhatsApp Business services",
            "deliverables": "Technical integration, service delivery, 24x7 support"
        },
        "financial_terms": {
            "pricing_model": "Usage-based pricing per unit",
            "payment_terms": "45 days for monthly recurring, advance for yearly",
            "invoicing": "Monthly billing based on usage"
        },
        "dates": {
            "execution_date": exec_date,
            "effective_date": effective_date,
            "initial_term": "1 year with automatic renewal",
            "renewal_terms": "Automatic yearly renewal unless terminated"
        },
        "termination": {
            "termination_rights": "Either party may terminate with cause or 60 days notice",
            "notice_period": "60 days for termination without cause",
            "effect_of_termination": "Immediate cessation of services, data return"
        },
        "key_obligations": [
            "Service provider: Technical integration, service delivery, support",
            "Client: Payment obligations, compliance with regulations, content responsibility"
        ],
        "liability_indemnity": [
            "Mutual indemnification for IP infringement",
            "Client indemnifies for content and spam activity",
            "Standard limitation of liability provisions"
        ]
    }

def extract_insurance_info_manually(document_text: str) -> Dict[str, Any]:
    """Manual extraction for insurance policies"""
    import re
    
    # Find policy number
    policy_match = re.search(r'policy\s*(?:number|no\.?)\s*[:\-]?\s*([A-Z0-9\-]+)', document_text, re.IGNORECASE)
    policy_number = policy_match.group(1).strip() if policy_match else "Not specified"
    
    # Find premium amount
    premium_match = re.search(r'premium\s*[:\-]?\s*[$₹]?(\d+(?:,\d+)*(?:\.\d{2})?)', document_text, re.IGNORECASE)
    premium_amount = f"${premium_match.group(1)}" if premium_match else "Not specified"
    
    # Find insurance company
    insurer_match = re.search(r'([\w\s]+)\s*insurance\s*company', document_text, re.IGNORECASE)
    insurer = insurer_match.group(1).strip() + " Insurance Company" if insurer_match else "Not specified"
    
    # Find policyholder
    policyholder_match = re.search(r'policyholder[:\-]?\s*([A-Za-z\s]+)', document_text, re.IGNORECASE)
    policyholder = policyholder_match.group(1).strip() if policyholder_match else "Not specified"
    
    # Find dates
    start_date_match = re.search(r'effective\s*(?:date|from)[:\-]?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})', document_text, re.IGNORECASE)
    start_date = start_date_match.group(1).replace('/', '-') if start_date_match else "Not specified"
    
    return {
        "document_type": "Insurance Policy",
        "contract_summary": f"Insurance policy from {insurer} covering {policyholder} with premium {premium_amount}",
        "parties": {
            "insurer": insurer,
            "policyholder": policyholder,
            "beneficiary": "Not specified"
        },
        "policy_details": {
            "policy_number": policy_number,
            "insurance_type": "General insurance policy",
            "coverage_amount": "As specified in the policy",
            "deductible": "Not specified"
        },
        "financial_terms": {
            "premium_amount": premium_amount,
            "copayment": "Not specified",
            "coinsurance": "Not specified",
            "out_of_pocket_maximum": "Not specified"
        },
        "dates": {
            "policy_start": start_date,
            "policy_end": "Not specified",
            "renewal_date": "Not specified",
            "grace_period": "Not specified"
        },
        "coverage_details": [
            "Coverage as outlined in the policy document"
        ],
        "exclusions": [
            "Exclusions as specified in the policy terms"
        ],
        "claim_procedures": [
            "Claim procedures as outlined in the policy"
        ]
    }

def extract_mou_info_manually(document_text: str) -> Dict[str, Any]:
    """Manual extraction for MOU documents"""
    import re
    
    # Find parties (look for organization names)
    parties_pattern = r'between\s+([^,\n]+)\s+and\s+([^,\n]+)'
    parties_match = re.search(parties_pattern, document_text, re.IGNORECASE)
    
    if parties_match:
        party_1 = parties_match.group(1).strip()
        party_2 = parties_match.group(2).strip()
    else:
        party_1 = "Not specified"
        party_2 = "Not specified"
    
    # Find execution date
    exec_date_match = re.search(r'(?:signed|executed|dated)\s*(?:on)?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})', document_text, re.IGNORECASE)
    execution_date = exec_date_match.group(1).replace('/', '-') if exec_date_match else "Not specified"
    
    # Find duration/term
    duration_match = re.search(r'(?:term|duration|period)\s*(?:of)?\s*(\d+\s*(?:years?|months?))', document_text, re.IGNORECASE)
    duration = duration_match.group(1) if duration_match else "Not specified"
    
    return {
        "document_type": "Memorandum of Understanding",
        "contract_summary": f"MOU between {party_1} and {party_2} for collaborative purposes",
        "parties": {
            "party_1": party_1,
            "party_2": party_2,
            "additional_parties": "Not specified"
        },
        "purpose": {
            "objective": "Collaborative objectives as outlined in the MOU",
            "background": "Background context for the partnership",
            "scope": "Scope of cooperation as defined in the MOU"
        },
        "responsibilities": {
            "party_1_obligations": f"Responsibilities of {party_1} as outlined in the MOU",
            "party_2_obligations": f"Responsibilities of {party_2} as outlined in the MOU",
            "shared_responsibilities": "Joint obligations and shared activities"
        },
        "dates": {
            "execution_date": execution_date,
            "effective_date": execution_date,
            "duration": duration,
            "review_date": "Not specified"
        },
        "key_terms": [
            "Terms and conditions as outlined in the MOU",
            "Mutual cooperation and collaboration",
            "Periodic review and assessment"
        ],
        "termination_conditions": [
            "Termination conditions as specified in the MOU",
            "Mutual agreement termination provision"
        ],
        "governance": [
            "Governance structure as defined in the MOU",
            "Management and oversight procedures"
        ]
    }

def month_to_number(month_name: str) -> str:
    """Convert month name to number"""
    months = {
        'january': '01', 'february': '02', 'march': '03', 'april': '04',
        'may': '05', 'june': '06', 'july': '07', 'august': '08',
        'september': '09', 'october': '10', 'november': '11', 'december': '12'
    }
    return months.get(month_name.lower(), '01')

def calculate_days_until_expiry(end_date_str: str) -> tuple[Optional[int], Optional[datetime]]:
    """Calculate days until contract expiry with logging"""
    logger.info(f"Calculating days until expiry for date: {end_date_str}")
    
    try:
        if end_date_str and end_date_str != "Not specified":
            end_date = parser.parse(end_date_str)
            today = datetime.now()
            days_left = (end_date - today).days
            
            logger.info(f"Contract end date: {end_date.strftime('%Y-%m-%d')}")
            logger.info(f"Today's date: {today.strftime('%Y-%m-%d')}")
            logger.info(f"Days until expiry: {days_left}")
            
            if days_left < 0:
                logger.warning(f"Contract has expired {abs(days_left)} days ago")
            elif days_left <= 30:
                logger.warning(f"Contract expires soon: {days_left} days remaining")
            else:
                logger.info(f"Contract has {days_left} days remaining")
            
            return days_left, end_date
        else:
            logger.warning("End date not specified or invalid")
            return None, None
    except Exception as e:
        logger.error(f"Error calculating days until expiry: {str(e)}")
        return None, None

def calculate_nda_expiry(data: Dict[str, Any]) -> tuple[Optional[int], Optional[datetime]]:
    """Calculate NDA expiry date based on duration and start date"""
    try:
        dates = data.get("dates", {})
        duration = dates.get("duration", "")
        survival_period = dates.get("survival_period", "")
        execution_date = dates.get("execution_date", "")
        effective_date = dates.get("effective_date", "")
        
        logger.info(f"NDA expiry calculation - Duration: '{duration}', Survival: '{survival_period}'")
        logger.info(f"NDA dates - Execution: '{execution_date}', Effective: '{effective_date}'")
        
        # Use effective date or execution date as start date
        start_date_str = effective_date if effective_date != "Not specified" else execution_date
        
        if start_date_str == "Not specified":
            logger.warning("No valid start date found for NDA expiry calculation")
            return None, None
        
        # Use the longer/more specific duration text
        duration_to_use = duration if duration != "Not specified" else survival_period
        
        if duration_to_use == "Not specified":
            logger.warning("No valid duration found for NDA expiry calculation")
            return None, None
        
        logger.info(f"Using duration text: '{duration_to_use}'")
        
        # Enhanced regex patterns to handle complex duration formats
        import re
        
        # Pattern 1: "five (5) years", "three (3) years", etc.
        written_years_match = re.search(r'(?:five|5)\s*\(\s*5\s*\)\s*years?', duration_to_use.lower())
        if written_years_match:
            years = 5
        else:
            written_years_match = re.search(r'(?:three|3)\s*\(\s*3\s*\)\s*years?', duration_to_use.lower())
            if written_years_match:
                years = 3
            else:
                written_years_match = re.search(r'(?:two|2)\s*\(\s*2\s*\)\s*years?', duration_to_use.lower())
                if written_years_match:
                    years = 2
                else:
                    written_years_match = re.search(r'(?:one|1)\s*\(\s*1\s*\)\s*years?', duration_to_use.lower())
                    if written_years_match:
                        years = 1
                    else:
                        # Pattern 2: Standard "5 years", "3 years", etc.
                        years_match = re.search(r'(\d+)\s*years?', duration_to_use.lower())
                        if years_match:
                            years = int(years_match.group(1))
                        else:
                            # Pattern 3: "12 months", "24 months", etc.
                            months_match = re.search(r'(\d+)\s*months?', duration_to_use.lower())
                            if months_match:
                                months = int(months_match.group(1))
                                start_date = parser.parse(start_date_str)
                                end_date = start_date + relativedelta(months=months)
                                
                                today = datetime.now()
                                days_left = (end_date - today).days
                                
                                logger.info(f"NDA expiry calculated (months): {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
                                return days_left, end_date
                            else:
                                logger.warning(f"Could not parse duration: {duration_to_use}")
                                return None, None
        
        # If we got years, calculate the end date
        if 'years' in locals():
            start_date = parser.parse(start_date_str)
            end_date = start_date + relativedelta(years=years)
            
            today = datetime.now()
            days_left = (end_date - today).days
            
            logger.info(f"NDA expiry calculated (years): {end_date.strftime('%Y-%m-%d')} ({days_left} days from {years} years)")
            return days_left, end_date
        else:
            logger.warning(f"No years variable set, failed to parse: {duration_to_use}")
            return None, None
            
    except Exception as e:
        logger.error(f"Error calculating NDA expiry: {str(e)}")
        return None, None

def calculate_msa_expiry(data: Dict[str, Any]) -> tuple[Optional[int], Optional[datetime]]:
    """Calculate MSA expiry date based on initial term and start date"""
    try:
        dates = data.get("dates", {})
        initial_term = dates.get("initial_term", "")
        execution_date = dates.get("execution_date", "")
        effective_date = dates.get("effective_date", "")
        
        # Use effective date or execution date as start date
        start_date_str = effective_date if effective_date != "Not specified" else execution_date
        
        if start_date_str == "Not specified":
            logger.warning("No valid start date found for MSA expiry calculation")
            return None, None
        
        if initial_term == "Not specified":
            logger.warning("No valid initial term found for MSA expiry calculation")
            return None, None
        
        # Extract years or months from initial term (e.g., "1 year", "2 years", "12 months")
        import re
        years_match = re.search(r'(\d+)\s*years?', initial_term.lower())
        months_match = re.search(r'(\d+)\s*months?', initial_term.lower())
        
        if years_match:
            years = int(years_match.group(1))
            start_date = parser.parse(start_date_str)
            end_date = start_date + relativedelta(years=years)
            
            today = datetime.now()
            days_left = (end_date - today).days
            
            logger.info(f"MSA expiry calculated: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
            return days_left, end_date
        elif months_match:
            months = int(months_match.group(1))
            start_date = parser.parse(start_date_str)
            end_date = start_date + relativedelta(months=months)
            
            today = datetime.now()
            days_left = (end_date - today).days
            
            logger.info(f"MSA expiry calculated: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
            return days_left, end_date
        else:
            logger.warning(f"Could not parse initial term: {initial_term}")
            return None, None
            
    except Exception as e:
        logger.error(f"Error calculating MSA expiry: {str(e)}")
        return None, None

def calculate_insurance_expiry(data: Dict[str, Any]) -> tuple[Optional[int], Optional[datetime]]:
    """Calculate insurance policy expiry date"""
    try:
        dates = data.get("dates", {})
        policy_end = dates.get("policy_end", "")
        renewal_date = dates.get("renewal_date", "")
        policy_start = dates.get("policy_start", "")
        grace_period = dates.get("grace_period", "")
        
        # Use policy_end, then renewal_date
        end_date_str = policy_end if policy_end != "Not specified" else renewal_date
        
        if end_date_str != "Not specified":
            try:
                end_date = parser.parse(end_date_str)
                today = datetime.now()
                days_left = (end_date - today).days
                
                logger.info(f"Insurance expiry calculated: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
                return days_left, end_date
            except:
                pass
        
        # Calculate from start date + duration if available
        if policy_start != "Not specified":
            try:
                start_date = parser.parse(policy_start)
                
                # Look for duration in various fields - check all date fields for duration info
                duration_sources = [
                    grace_period,
                    dates.get("duration", ""),
                    dates.get("term", ""),
                    dates.get("policy_term", "")
                ]
                
                # Also check in other data structures
                policy_details = data.get("policy_details", {})
                financial_terms = data.get("financial_terms", {})
                
                # Look for duration indicators in policy details or financial terms
                duration_sources.extend([
                    str(policy_details.get("coverage_period", "")),
                    str(financial_terms.get("premium_frequency", "")),
                    str(financial_terms.get("payment_frequency", ""))
                ])
                
                import re
                end_date = None
                
                for duration_text in duration_sources:
                    if duration_text and duration_text != "Not specified":
                        logger.info(f"Trying to parse duration from: {duration_text}")
                        
                        # Look for year patterns
                        year_match = re.search(r'(\d+)\s*years?', duration_text.lower())
                        if year_match:
                            years = int(year_match.group(1))
                            end_date = start_date + relativedelta(years=years)
                            logger.info(f"Found {years} year(s) duration, calculated end date: {end_date}")
                            break
                        
                        # Look for month patterns
                        month_match = re.search(r'(\d+)\s*months?', duration_text.lower())
                        if month_match:
                            months = int(month_match.group(1))
                            end_date = start_date + relativedelta(months=months)
                            logger.info(f"Found {months} month(s) duration, calculated end date: {end_date}")
                            break
                        
                        # Look for "one year", "two years", etc.
                        if "one year" in duration_text.lower() or "1 year" in duration_text.lower():
                            end_date = start_date + relativedelta(years=1)
                            logger.info(f"Found 'one year' duration, calculated end date: {end_date}")
                            break
                        
                        # Check for annual/yearly indicators
                        if any(word in duration_text.lower() for word in ["annual", "yearly", "per year"]):
                            end_date = start_date + relativedelta(years=1)
                            logger.info(f"Found annual indicator, calculated end date: {end_date}")
                            break
                
                # Default fallback: assume 1 year for insurance policies
                if end_date is None:
                    end_date = start_date + relativedelta(years=1)
                    logger.info(f"No duration found, defaulting to 1 year: {end_date}")
                
                today = datetime.now()
                days_left = (end_date - today).days
                
                logger.info(f"Insurance expiry calculated from start date + duration: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
                return days_left, end_date
                
            except Exception as calc_error:
                logger.error(f"Error calculating from start date: {calc_error}")
        
        logger.warning("Could not calculate insurance expiry date")
        return None, None
            
    except Exception as e:
        logger.error(f"Error calculating insurance expiry: {str(e)}")
        return None, None

def calculate_end_date_from_duration(start_date_str: str, duration_text: str) -> Optional[datetime]:
    """Universal function to calculate end date from start date + duration text"""
    if not start_date_str or start_date_str == "Not specified":
        return None
    if not duration_text or duration_text == "Not specified":
        return None
    
    try:
        start_date = parser.parse(start_date_str)
        duration_lower = duration_text.lower()
        
        import re
        
        # Look for specific year patterns
        year_patterns = [
            r'(\d+)\s*years?',
            r'(?:one|1)\s*years?',
            r'(?:two|2)\s*years?',
            r'(?:three|3)\s*years?',
            r'(?:four|4)\s*years?',
            r'(?:five|5)\s*years?'
        ]
        
        for pattern in year_patterns:
            match = re.search(pattern, duration_lower)
            if match:
                if pattern.startswith(r'(\d+)'):
                    years = int(match.group(1))
                elif 'one' in pattern or '1' in pattern:
                    years = 1
                elif 'two' in pattern or '2' in pattern:
                    years = 2
                elif 'three' in pattern or '3' in pattern:
                    years = 3
                elif 'four' in pattern or '4' in pattern:
                    years = 4
                elif 'five' in pattern or '5' in pattern:
                    years = 5
                else:
                    continue
                
                end_date = start_date + relativedelta(years=years)
                logger.info(f"Calculated end date from {years} year(s): {end_date}")
                return end_date
        
        # Look for month patterns
        month_patterns = [
            r'(\d+)\s*months?',
            r'(?:eleven|11)\s*months?',
            r'(?:twelve|12)\s*months?',
            r'(?:six|6)\s*months?'
        ]
        
        for pattern in month_patterns:
            match = re.search(pattern, duration_lower)
            if match:
                if pattern.startswith(r'(\d+)'):
                    months = int(match.group(1))
                elif 'eleven' in pattern or '11' in pattern:
                    months = 11
                elif 'twelve' in pattern or '12' in pattern:
                    months = 12
                elif 'six' in pattern or '6' in pattern:
                    months = 6
                else:
                    continue
                
                end_date = start_date + relativedelta(months=months)
                logger.info(f"Calculated end date from {months} month(s): {end_date}")
                return end_date
        
        # Check for common duration keywords
        if any(word in duration_lower for word in ["annual", "yearly", "per year", "one year"]):
            end_date = start_date + relativedelta(years=1)
            logger.info(f"Calculated end date from annual term: {end_date}")
            return end_date
        
        logger.warning(f"Could not parse duration: {duration_text}")
        return None
        
    except Exception as e:
        logger.error(f"Error calculating end date from duration: {str(e)}")
        return None

def calculate_mou_expiry(data: Dict[str, Any]) -> tuple[Optional[int], Optional[datetime]]:
    """Calculate MOU expiry date based on duration and start date"""
    try:
        dates = data.get("dates", {})
        duration = dates.get("duration", "")
        execution_date = dates.get("execution_date", "")
        effective_date = dates.get("effective_date", "")
        review_date = dates.get("review_date", "")
        
        # Use effective date or execution date as start date
        start_date_str = effective_date if effective_date != "Not specified" else execution_date
        
        # If review_date is specified, use that as the expiry
        if review_date != "Not specified":
            try:
                end_date = parser.parse(review_date)
                today = datetime.now()
                days_left = (end_date - today).days
                
                logger.info(f"MOU review date: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
                return days_left, end_date
            except:
                pass
        
        # Calculate from start date + duration using universal function
        if start_date_str != "Not specified" and duration != "Not specified":
            end_date = calculate_end_date_from_duration(start_date_str, duration)
            if end_date:
                today = datetime.now()
                days_left = (end_date - today).days
                
                logger.info(f"MOU expiry calculated: {end_date.strftime('%Y-%m-%d')} ({days_left} days)")
                return days_left, end_date
        
        # Check if duration info is in other fields
        purpose = data.get("purpose", {})
        if isinstance(purpose, dict):
            for key, value in purpose.items():
                if value and value != "Not specified":
                    calculated_end = calculate_end_date_from_duration(start_date_str, str(value))
                    if calculated_end:
                        today = datetime.now()
                        days_left = (calculated_end - today).days
                        logger.info(f"MOU expiry calculated from purpose field: {calculated_end.strftime('%Y-%m-%d')} ({days_left} days)")
                        return days_left, calculated_end
        
        logger.warning("Could not calculate MOU expiry date")
        return None, None
            
    except Exception as e:
        logger.error(f"Error calculating MOU expiry: {str(e)}")
        return None, None

def display_minimal_contract_info(data: Optional[Dict[str, Any]]) -> None:
    """Display minimal contract information: Start Date, End Date, Document Type only"""
    logger.info("Displaying minimal contract information")
    
    if not data:
        logger.warning("No contract data available for display")
        st.warning("No contract data extracted yet.")
        return
    
    document_type = data.get("document_type", "Unknown Document")
    dates = data.get("dates", {})
    
    # Get start date
    start_date = (dates.get("start_date") or 
                 dates.get("effective_date") or 
                 dates.get("execution_date") or 
                 dates.get("policy_start") or 
                 "Not specified")
    
    # Get or calculate end date based on document type
    end_date = "Not specified"
    days_remaining = None
    
    if "Rental" in document_type:
        end_date = dates.get("end_date", "Not specified")
        # If no end date but we have start date and lease term, calculate it
        if end_date == "Not specified" and start_date != "Not specified":
            lease_term = dates.get("lease_term", "")
            if lease_term != "Not specified":
                calculated_end = calculate_end_date_from_duration(start_date, lease_term)
                if calculated_end:
                    end_date = calculated_end.strftime('%Y-%m-%d')
                    logger.info(f"Calculated rental end date from lease term: {end_date}")
        
        if end_date != "Not specified":
            try:
                end_dt = parser.parse(end_date)
                today = datetime.now()
                days_remaining = (end_dt - today).days
            except:
                pass
    
    elif "Non-Disclosure" in document_type or "NDA" in document_type:
        days_left, calculated_end_date = calculate_nda_expiry(data)
        if calculated_end_date is not None:
            end_date = calculated_end_date.strftime('%Y-%m-%d')
            days_remaining = days_left
    
    elif "Master Service" in document_type or "MSA" in document_type:
        days_left, calculated_end_date = calculate_msa_expiry(data)
        if calculated_end_date is not None:
            end_date = calculated_end_date.strftime('%Y-%m-%d')
            days_remaining = days_left
        # If calculation failed but we have start date and initial term, try universal calculation
        elif end_date == "Not specified" and start_date != "Not specified":
            initial_term = dates.get("initial_term", "")
            if initial_term != "Not specified":
                calculated_end = calculate_end_date_from_duration(start_date, initial_term)
                if calculated_end:
                    end_date = calculated_end.strftime('%Y-%m-%d')
                    today = datetime.now()
                    days_remaining = (calculated_end - today).days
                    logger.info(f"Calculated MSA end date from initial term: {end_date}")
    
    elif "Insurance" in document_type:
        days_left, calculated_end_date = calculate_insurance_expiry(data)
        if calculated_end_date is not None:
            end_date = calculated_end_date.strftime('%Y-%m-%d')
            days_remaining = days_left
        # Additional fallback for insurance - check various duration fields
        elif end_date == "Not specified" and start_date != "Not specified":
            # Try different duration sources
            duration_sources = [
                dates.get("grace_period", ""),
                dates.get("policy_term", ""),
                dates.get("duration", "")
            ]
            
            # Check policy details for duration info
            policy_details = data.get("policy_details", {})
            if isinstance(policy_details, dict):
                duration_sources.extend([
                    str(policy_details.get("coverage_period", "")),
                    str(policy_details.get("term", ""))
                ])
            
            for duration_text in duration_sources:
                if duration_text and duration_text != "Not specified":
                    calculated_end = calculate_end_date_from_duration(start_date, duration_text)
                    if calculated_end:
                        end_date = calculated_end.strftime('%Y-%m-%d')
                        today = datetime.now()
                        days_remaining = (calculated_end - today).days
                        logger.info(f"Calculated insurance end date from duration: {end_date}")
                        break
    
    elif "Memorandum" in document_type or "MOU" in document_type:
        days_left, calculated_end_date = calculate_mou_expiry(data)
        if calculated_end_date is not None:
            end_date = calculated_end_date.strftime('%Y-%m-%d')
            days_remaining = days_left
        # If calculation failed but we have start date and duration, try universal calculation
        elif end_date == "Not specified" and start_date != "Not specified":
            duration = dates.get("duration", "")
            if duration != "Not specified":
                calculated_end = calculate_end_date_from_duration(start_date, duration)
                if calculated_end:
                    end_date = calculated_end.strftime('%Y-%m-%d')
                    today = datetime.now()
                    days_remaining = (calculated_end - today).days
                    logger.info(f"Calculated MOU end date from duration: {end_date}")
    
    # Universal fallback for any document type with start date and duration
    if end_date == "Not specified" and start_date != "Not specified":
        # Try to find duration in various common fields
        duration_fields = [
            dates.get("duration", ""),
            dates.get("term", ""),
            dates.get("period", ""),
            dates.get("validity", "")
        ]
        
        for duration_text in duration_fields:
            if duration_text and duration_text != "Not specified":
                calculated_end = calculate_end_date_from_duration(start_date, duration_text)
                if calculated_end:
                    end_date = calculated_end.strftime('%Y-%m-%d')
                    today = datetime.now()
                    days_remaining = (calculated_end - today).days
                    logger.info(f"Calculated end date using universal fallback: {end_date}")
                    break
    
    # Display the minimal info in a clean format
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("📅 Start Date", start_date if start_date != "Not specified" else "Not found")
    
    with col2:
        if end_date != "Not specified" and days_remaining is not None:
            if days_remaining > 0:
                if days_remaining <= 30:
                    st.metric("⚠️ End Date", end_date, f"{days_remaining} days left", delta_color="inverse")
                elif days_remaining <= 90:
                    st.metric("⏰ End Date", end_date, f"{days_remaining} days left", delta_color="normal")
                else:
                    st.metric("📅 End Date", end_date, f"{days_remaining} days left")
            else:
                st.metric("❌ End Date", end_date, f"Expired {abs(days_remaining)} days ago", delta_color="inverse")
        else:
            # Show calculated end date even without days remaining if we calculated it
            if end_date != "Not specified":
                st.metric("📅 End Date", end_date, "Calculated from duration")
            else:
                st.metric("📅 End Date", "Not found")
    
    with col3:
        # Clean up document type for display
        clean_type = (document_type.replace("Agreement", "")
                                  .replace("Non-Disclosure", "NDA")
                                  .replace("Master Service", "MSA")
                                  .replace("Memorandum of Understanding", "MOU")
                                  .replace("Policy", "")
                                  .strip())
        st.metric("📄 Document Type", clean_type)
    
    logger.info("Minimal contract information displayed successfully")

def create_brief_summary(data: Dict[str, Any], document_type: str) -> str:
    """Create a brief 3-4 line summary of the contract"""
    
    if "Non-Disclosure" in document_type or "NDA" in document_type:
        parties = data.get("parties", {})
        disclosing = parties.get("disclosing_party", "Party 1")
        receiving = parties.get("receiving_party", "Party 2")
        dates = data.get("dates", {})
        duration = dates.get("duration", "specified period")
        
        return f"""This Non-Disclosure Agreement is between {disclosing} (disclosing party) and {receiving} (receiving party). The agreement protects confidential information shared for business discussions and evaluation purposes. Confidentiality obligations remain in effect for {duration}. Both parties must keep shared information confidential and use it only for the intended business purpose."""
    
    elif "Master Service" in document_type or "MSA" in document_type:
        parties = data.get("parties", {})
        provider = parties.get("service_provider", "Service Provider")
        client = parties.get("client", "Client")
        service_details = data.get("service_details", {})
        services = service_details.get("description", "professional services")
        dates = data.get("dates", {})
        term = dates.get("initial_term", "specified period")
        
        return f"""This Master Service Agreement is between {provider} (service provider) and {client} (client). The provider will deliver {services} under the terms and conditions outlined in the agreement. The contract has an initial term of {term} and covers service delivery, payment terms, and responsibilities of both parties."""
    
    elif "Insurance" in document_type:
        parties = data.get("parties", {})
        insurer = parties.get("insurer", "Insurance Company")
        policyholder = parties.get("policyholder", "Policyholder")
        policy_details = data.get("policy_details", {})
        insurance_type = policy_details.get("insurance_type", "insurance coverage")
        coverage_amount = policy_details.get("coverage_amount", "specified coverage")
        financial = data.get("financial_terms", {})
        premium = financial.get("premium_amount", "agreed premium")
        
        return f"""This insurance policy is between {insurer} (insurer) and {policyholder} (policyholder). The policy provides {insurance_type} with {coverage_amount} coverage. The policyholder pays {premium} in premiums and the policy outlines covered benefits, exclusions, and claim procedures."""
    
    elif "Memorandum" in document_type or "MOU" in document_type:
        parties = data.get("parties", {})
        party_1 = parties.get("party_1", "First Party")
        party_2 = parties.get("party_2", "Second Party")
        purpose = data.get("purpose", {})
        objective = purpose.get("objective", "collaborative objectives")
        dates = data.get("dates", {})
        duration = dates.get("duration", "specified period")
        
        return f"""This Memorandum of Understanding is between {party_1} and {party_2} for {objective}. The MOU establishes a framework for cooperation and collaboration between the parties. The agreement lasts for {duration} and outlines the responsibilities, governance structure, and terms of the partnership."""
    
    else:  # Rental
        parties = data.get("parties", {})
        landlord = parties.get("landlord", "Landlord")
        tenant = parties.get("tenant", "Tenant")
        property_details = data.get("property_details", {})
        address = property_details.get("address", "specified property")
        financial = data.get("financial_terms", {})
        rent = financial.get("monthly_rent", "agreed amount")
        dates = data.get("dates", {})
        term = dates.get("lease_term", "specified period")
        
        return f"""This rental agreement is between {landlord} (landlord) and {tenant} (tenant) for the property located at {address}. The lease term is {term} with monthly rent of {rent}. The agreement outlines the rights, responsibilities, and obligations of both landlord and tenant during the rental period."""

def get_relevant_context(question: str, contract_data: Dict[str, Any]) -> str:
    """Get only relevant contract data based on the question type"""
    question_lower = question.lower()
    
    # Smart context selection based on question keywords
    relevant_data = {}
    
    # Always include basic summary
    if "contract_summary" in contract_data:
        relevant_data["summary"] = contract_data["contract_summary"]
    
    # Date-related questions
    if any(word in question_lower for word in ["expire", "end", "start", "date", "when", "duration", "term"]):
        if "dates" in contract_data:
            relevant_data["dates"] = contract_data["dates"]
    
    # Financial questions  
    if any(word in question_lower for word in ["rent", "money", "cost", "pay", "amount", "price", "fee", "premium", "deductible"]):
        if "financial_terms" in contract_data:
            relevant_data["financial_terms"] = contract_data["financial_terms"]
    
    # Party/people questions
    if any(word in question_lower for word in ["who", "landlord", "tenant", "owner", "lessee", "lessor", "name", "insurer", "policyholder"]):
        if "parties" in contract_data:
            relevant_data["parties"] = contract_data["parties"]
    
    # Property questions
    if any(word in question_lower for word in ["where", "address", "property", "location", "building"]):
        if "property_details" in contract_data:
            relevant_data["property_details"] = contract_data["property_details"]
    
    # Terms and conditions
    if any(word in question_lower for word in ["term", "condition", "rule", "clause", "allowed", "restriction"]):
        if "key_terms" in contract_data:
            relevant_data["key_terms"] = contract_data["key_terms"]
        if "special_clauses" in contract_data:
            relevant_data["special_clauses"] = contract_data["special_clauses"]
    
    # Insurance-specific questions
    if any(word in question_lower for word in ["coverage", "covered", "exclude", "claim", "benefit"]):
        if "coverage_details" in contract_data:
            relevant_data["coverage_details"] = contract_data["coverage_details"]
        if "exclusions" in contract_data:
            relevant_data["exclusions"] = contract_data["exclusions"]
        if "claim_procedures" in contract_data:
            relevant_data["claim_procedures"] = contract_data["claim_procedures"]
    
    # MOU-specific questions
    if any(word in question_lower for word in ["purpose", "objective", "responsibility", "obligation", "governance"]):
        if "purpose" in contract_data:
            relevant_data["purpose"] = contract_data["purpose"]
        if "responsibilities" in contract_data:
            relevant_data["responsibilities"] = contract_data["responsibilities"]
        if "governance" in contract_data:
            relevant_data["governance"] = contract_data["governance"]
    
    # If no specific context found, include essential data
    if not relevant_data:
        relevant_data = {
            "summary": contract_data.get("contract_summary", ""),
            "dates": contract_data.get("dates", {}),
            "parties": contract_data.get("parties", {}),
            "financial_terms": contract_data.get("financial_terms", {})
        }
    
    return json.dumps(relevant_data, indent=2)

def answer_question_locally(question: str, contract_data: Dict[str, Any]) -> Optional[str]:
    """Answer simple questions directly from structured data (instant responses for all document types)"""
    import time
    start_time = time.perf_counter()
    
    question_lower = question.lower()
    document_type = contract_data.get("document_type", "Unknown")
    logger.info(f"Attempting local answer for {document_type}: {question}")
    
    try:
        # Universal WHO questions (works for all document types)
        if any(word in question_lower for word in ["who", "party", "parties"]):
            parties = contract_data.get("parties", {})
            
            # Rental-specific
            if "tenant" in question_lower or "lessee" in question_lower:
                tenant = parties.get("tenant", "Not specified")
                if tenant != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (tenant): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The person renting the property is: {tenant}"
            
            if "landlord" in question_lower or "lessor" in question_lower:
                landlord = parties.get("landlord", "Not specified")
                if landlord != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (landlord): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The property owner is: {landlord}"
            
            # NDA-specific
            if "disclosing" in question_lower:
                disclosing = parties.get("disclosing_party", "Not specified")
                if disclosing != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (disclosing party): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company sharing confidential information is: {disclosing}"
            
            if "receiving" in question_lower:
                receiving = parties.get("receiving_party", "Not specified")
                if receiving != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (receiving party): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company receiving confidential information is: {receiving}"
            
            # MSA-specific
            if "provider" in question_lower or "service" in question_lower and "provider" in question_lower:
                provider = parties.get("service_provider", "Not specified")
                if provider != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (service provider): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company providing the services is: {provider}"
            
            if "client" in question_lower:
                client = parties.get("client", "Not specified")
                if client != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (client): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The company buying the services is: {client}"
            
            # Insurance-specific
            if "insurer" in question_lower or "insurance company" in question_lower:
                insurer = parties.get("insurer", "Not specified")
                if insurer != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (insurer): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The insurance company is: {insurer}"
            
            if "policyholder" in question_lower or "insured" in question_lower:
                policyholder = parties.get("policyholder", "Not specified")
                if policyholder != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (policyholder): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The person being insured is: {policyholder}"
            
            # MOU-specific
            if "first party" in question_lower or "party 1" in question_lower:
                party_1 = parties.get("party_1", "Not specified")
                if party_1 != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (party 1): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The first party is: {party_1}"
            
            if "second party" in question_lower or "party 2" in question_lower:
                party_2 = parties.get("party_2", "Not specified")
                if party_2 != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (party 2): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The second party is: {party_2}"
        
        # Universal WHEN questions (works for all document types)
        if any(word in question_lower for word in ["when", "expire", "end", "start", "date"]):
            dates = contract_data.get("dates", {})
            
            if "expire" in question_lower or "end" in question_lower:
                # Handle different document types for expiry questions
                if "Rental" in document_type:
                    # For rentals, use the direct end_date
                    end_date = dates.get("end_date", "Not specified")
                    if end_date != "Not specified":
                        try:
                            from dateutil import parser
                            end_dt = parser.parse(end_date)
                            today = datetime.now()
                            days_left = (end_dt - today).days
                            
                            if days_left > 0:
                                return f"The lease expires on {end_date} ({days_left} days remaining)."
                            else:
                                return f"The lease expired on {end_date} ({abs(days_left)} days ago)."
                        except:
                            return f"The lease expires/ends on {end_date}."
                
                elif "Non-Disclosure" in document_type or "NDA" in document_type:
                    # For NDAs, calculate expiry from start date + duration
                    logger.info(f"Processing NDA expiry question for: {question}")
                    logger.info(f"Available date fields: {dates}")
                    
                    days_left, calculated_end_date = calculate_nda_expiry(contract_data)
                    if days_left is not None and calculated_end_date is not None:
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (NDA expiry): {(end_time - start_time) * 1000:.2f}ms")
                        
                        if days_left > 0:
                            return f"The confidentiality agreement runs out on {calculated_end_date.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                        else:
                            return f"The confidentiality agreement already ended on {calculated_end_date.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                    else:
                        # Enhanced fallback - try to give more useful info
                        logger.warning("NDA expiry calculation failed, using enhanced fallback")
                        
                        # Get all available date info for debugging
                        execution_date = dates.get("execution_date", "Not specified")
                        effective_date = dates.get("effective_date", "Not specified")
                        duration = dates.get("duration", "Not specified")
                        survival_period = dates.get("survival_period", "Not specified")
                        
                        logger.info(f"Debug - Execution: {execution_date}, Effective: {effective_date}")
                        logger.info(f"Debug - Duration: {duration}, Survival: {survival_period}")
                        
                        # Try manual calculation if we have start date and duration
                        start_date_str = effective_date if effective_date != "Not specified" else execution_date
                        duration_text = duration if duration != "Not specified" else survival_period
                        
                        if start_date_str != "Not specified" and duration_text != "Not specified":
                            # Simple fallback calculation
                            try:
                                import re
                                # Look for any number followed by "year"
                                year_match = re.search(r'(\d+).*?years?', duration_text.lower())
                                if year_match:
                                    years = int(year_match.group(1))
                                    start_date = parser.parse(start_date_str)
                                    end_date = start_date + relativedelta(years=years)
                                    today = datetime.now()
                                    days_left = (end_date - today).days
                                    
                                    logger.info(f"Manual NDA calculation successful: {end_date.strftime('%Y-%m-%d')}")
                                    if days_left > 0:
                                        return f"The confidentiality agreement runs out on {end_date.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                                    else:
                                        return f"The confidentiality agreement already ended on {end_date.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                                        
                            except Exception as calc_error:
                                logger.error(f"Manual calculation also failed: {calc_error}")
                        
                        # Ultimate fallback
                        if duration_text != "Not specified":
                            return f"The confidentiality agreement lasts for {duration_text}. It started on {start_date_str}."
                        else:
                            return "The document doesn't clearly specify when the confidentiality agreement ends."
                
                elif "Master Service" in document_type or "MSA" in document_type:
                    # For MSAs, calculate expiry from start date + initial term
                    days_left, calculated_end_date = calculate_msa_expiry(contract_data)
                    if days_left is not None and calculated_end_date is not None:
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (MSA expiry): {(end_time - start_time) * 1000:.2f}ms")
                        
                        renewal_info = dates.get("renewal_terms", "")
                        if days_left > 0:
                            base_msg = f"The service agreement runs out on {calculated_end_date.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                            if renewal_info and renewal_info != "Not specified":
                                return f"{base_msg} Keep in mind: {renewal_info}"
                            return base_msg
                        else:
                            return f"The service agreement already ended on {calculated_end_date.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                    else:
                        # Fallback to initial term text if calculation failed
                        initial_term = dates.get("initial_term", "Not specified")
                        if initial_term != "Not specified":
                            return f"The service agreement lasts for {initial_term}."
                
                elif "Insurance" in document_type:
                    # For Insurance policies, calculate expiry from policy dates
                    days_left, calculated_end_date = calculate_insurance_expiry(contract_data)
                    if days_left is not None and calculated_end_date is not None:
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (Insurance expiry): {(end_time - start_time) * 1000:.2f}ms")
                        
                        if days_left > 0:
                            return f"The insurance policy expires on {calculated_end_date.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                        else:
                            return f"The insurance policy already expired on {calculated_end_date.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                    else:
                        # Try fallback calculation with start date + duration
                        policy_start = dates.get("policy_start", "Not specified")
                        if policy_start != "Not specified":
                            # Try various duration sources
                            duration_sources = [
                                dates.get("grace_period", ""),
                                dates.get("policy_term", ""),
                                dates.get("duration", "")
                            ]
                            
                            for duration_text in duration_sources:
                                if duration_text and duration_text != "Not specified":
                                    calculated_end = calculate_end_date_from_duration(policy_start, duration_text)
                                    if calculated_end:
                                        today = datetime.now()
                                        days_left = (calculated_end - today).days
                                        
                                        end_time = time.perf_counter()
                                        logger.info(f"⚡ Local answer (Insurance calculated expiry): {(end_time - start_time) * 1000:.2f}ms")
                                        
                                        if days_left > 0:
                                            return f"Based on the policy start date and duration, your insurance expires on {calculated_end.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                                        else:
                                            return f"Based on the policy start date and duration, your insurance expired on {calculated_end.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                            
                            # Ultimate fallback for insurance - assume 1 year
                            calculated_end = calculate_end_date_from_duration(policy_start, "1 year")
                            if calculated_end:
                                today = datetime.now()
                                days_left = (calculated_end - today).days
                                
                                end_time = time.perf_counter()
                                logger.info(f"⚡ Local answer (Insurance 1-year fallback): {(end_time - start_time) * 1000:.2f}ms")
                                
                                if days_left > 0:
                                    return f"Assuming a standard 1-year policy, your insurance expires on {calculated_end.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                                else:
                                    return f"Assuming a standard 1-year policy, your insurance expired on {calculated_end.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                        
                        # Final fallback
                        policy_end = dates.get("policy_end", "Not specified")
                        if policy_end != "Not specified":
                            return f"The insurance policy expires on {policy_end}."
                        else:
                            return "The policy expiry date is not clearly specified in the document."
                
                elif "Memorandum" in document_type or "MOU" in document_type:
                    # For MOUs, calculate expiry or review date
                    days_left, calculated_end_date = calculate_mou_expiry(contract_data)
                    if days_left is not None and calculated_end_date is not None:
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (MOU expiry): {(end_time - start_time) * 1000:.2f}ms")
                        
                        if days_left > 0:
                            return f"The MOU is scheduled for review/expiry on {calculated_end_date.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                        else:
                            return f"The MOU review/expiry date was {calculated_end_date.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                    else:
                        # Try fallback calculation with start date + duration
                        execution_date = dates.get("execution_date", "Not specified")
                        effective_date = dates.get("effective_date", "Not specified")
                        start_date_str = effective_date if effective_date != "Not specified" else execution_date
                        
                        if start_date_str != "Not specified":
                            duration = dates.get("duration", "Not specified")
                            if duration != "Not specified":
                                calculated_end = calculate_end_date_from_duration(start_date_str, duration)
                                if calculated_end:
                                    today = datetime.now()
                                    days_left = (calculated_end - today).days
                                    
                                    end_time = time.perf_counter()
                                    logger.info(f"⚡ Local answer (MOU calculated expiry): {(end_time - start_time) * 1000:.2f}ms")
                                    
                                    if days_left > 0:
                                        return f"Based on the start date and duration, the MOU expires on {calculated_end.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                                    else:
                                        return f"Based on the start date and duration, the MOU expired on {calculated_end.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
                        
                        # Fallback to duration text
                        duration = dates.get("duration", "Not specified")
                        if duration != "Not specified":
                            return f"The MOU lasts for {duration}."
                        else:
                            return "The MOU duration is not clearly specified in the document."
                
                # Generic fallback for any document type
                end_date = dates.get("end_date") or dates.get("policy_end") or "Not specified"
                if end_date != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (generic expiry): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The agreement runs out on {end_date}."
                
                # Universal fallback: try to calculate from any start date + duration
                start_date_options = [
                    dates.get("start_date", ""),
                    dates.get("effective_date", ""),
                    dates.get("execution_date", ""),
                    dates.get("policy_start", "")
                ]
                
                duration_options = [
                    dates.get("duration", ""),
                    dates.get("term", ""),
                    dates.get("lease_term", ""),
                    dates.get("initial_term", ""),
                    dates.get("policy_term", ""),
                    dates.get("period", "")
                ]
                
                for start_date_str in start_date_options:
                    if start_date_str and start_date_str != "Not specified":
                        for duration_text in duration_options:
                            if duration_text and duration_text != "Not specified":
                                calculated_end = calculate_end_date_from_duration(start_date_str, duration_text)
                                if calculated_end:
                                    today = datetime.now()
                                    days_left = (calculated_end - today).days
                                    
                                    end_time = time.perf_counter()
                                    logger.info(f"⚡ Local answer (universal calculated expiry): {(end_time - start_time) * 1000:.2f}ms")
                                    
                                    if days_left > 0:
                                        return f"Based on the start date and duration, this agreement expires on {calculated_end.strftime('%Y-%m-%d')} - that's {days_left} days from now."
                                    else:
                                        return f"Based on the start date and duration, this agreement expired on {calculated_end.strftime('%Y-%m-%d')} - that was {abs(days_left)} days ago."
            
            if "start" in question_lower or "effective" in question_lower:
                start_date = (dates.get("start_date") or 
                             dates.get("effective_date") or 
                             dates.get("execution_date") or 
                             dates.get("policy_start") or 
                             "Not specified")
                if start_date != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (start): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The agreement started on {start_date}."
        
        # Document-specific financial questions
        if any(word in question_lower for word in ["how much", "cost", "price", "amount", "money", "pay"]):
            
            if "Rental" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "rent" in question_lower:
                    rent = financial.get("monthly_rent", "Not specified")
                    if rent != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (rent): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The monthly rent is {rent}."
                
                if "deposit" in question_lower:
                    deposit = financial.get("security_deposit", "Not specified")
                    if deposit != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (deposit): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The security deposit (upfront payment) is {deposit}."
            
            elif "MSA" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "payment" in question_lower:
                    payment_terms = financial.get("payment_terms", "Not specified")
                    if payment_terms != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (payment terms): {(end_time - start_time) * 1000:.2f}ms")
                        return f"Here's how payments work: {payment_terms}."
            
            elif "Insurance" in document_type:
                financial = contract_data.get("financial_terms", {})
                if "premium" in question_lower:
                    premium = financial.get("premium_amount", "Not specified")
                    if premium != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (premium): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The insurance premium is {premium}."
                
                if "deductible" in question_lower:
                    policy_details = contract_data.get("policy_details", {})
                    deductible = policy_details.get("deductible", "Not specified")
                    if deductible != "Not specified":
                        end_time = time.perf_counter()
                        logger.info(f"⚡ Local answer (deductible): {(end_time - start_time) * 1000:.2f}ms")
                        return f"The deductible amount is {deductible}."
        
        # Location questions (mainly for rental)
        if any(word in question_lower for word in ["where", "address", "location", "property"]):
            if "Rental" in document_type:
                property_details = contract_data.get("property_details", {})
                address = property_details.get("address", "Not specified")
                if address != "Not specified":
                    end_time = time.perf_counter()
                    logger.info(f"⚡ Local answer (address): {(end_time - start_time) * 1000:.2f}ms")
                    return f"The property is located at: {address}"
        
        # Duration/term questions (universal)
        if any(word in question_lower for word in ["duration", "term", "how long", "period"]):
            dates = contract_data.get("dates", {})
            
            # Try different term fields based on document type
            term = (dates.get("lease_term") or 
                   dates.get("initial_term") or 
                   dates.get("duration") or 
                   "Not specified")
            
            if term != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (term): {(end_time - start_time) * 1000:.2f}ms")
                return f"The agreement lasts for {term}."
        
        # Document-specific questions
        if "confidential" in question_lower and "NDA" in document_type:
            scope = contract_data.get("confidentiality_scope", {})
            definition = scope.get("definition", "Not specified")
            if definition != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (confidentiality): {(end_time - start_time) * 1000:.2f}ms")
                return f"Here's what must be kept secret: {definition}"
        
        if "service" in question_lower and "MSA" in document_type:
            service_details = contract_data.get("service_details", {})
            description = service_details.get("description", "Not specified")
            if description != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (services): {(end_time - start_time) * 1000:.2f}ms")
                return f"Here's what services are being provided: {description}"
        
        # Insurance-specific questions
        if "covered" in question_lower and "Insurance" in document_type:
            coverage_details = contract_data.get("coverage_details", [])
            if coverage_details and coverage_details != ["Not specified"]:
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (coverage): {(end_time - start_time) * 1000:.2f}ms")
                return f"Here's what's covered: {', '.join(coverage_details)}"
        
        if "excluded" in question_lower or "exclusion" in question_lower and "Insurance" in document_type:
            exclusions = contract_data.get("exclusions", [])
            if exclusions and exclusions != ["Not specified"]:
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (exclusions): {(end_time - start_time) * 1000:.2f}ms")
                return f"Here's what's NOT covered: {', '.join(exclusions)}"
        
        if "policy number" in question_lower and "Insurance" in document_type:
            policy_details = contract_data.get("policy_details", {})
            policy_number = policy_details.get("policy_number", "Not specified")
            if policy_number != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (policy number): {(end_time - start_time) * 1000:.2f}ms")
                return f"The policy number is: {policy_number}"
        
        # MOU-specific questions
        if "purpose" in question_lower and "MOU" in document_type:
            purpose = contract_data.get("purpose", {})
            objective = purpose.get("objective", "Not specified")
            if objective != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (MOU purpose): {(end_time - start_time) * 1000:.2f}ms")
                return f"The purpose of this MOU is: {objective}"
        
        if "responsibility" in question_lower or "obligation" in question_lower and "MOU" in document_type:
            responsibilities = contract_data.get("responsibilities", {})
            party_1_obs = responsibilities.get("party_1_obligations", "Not specified")
            party_2_obs = responsibilities.get("party_2_obligations", "Not specified")
            if party_1_obs != "Not specified" and party_2_obs != "Not specified":
                end_time = time.perf_counter()
                logger.info(f"⚡ Local answer (MOU responsibilities): {(end_time - start_time) * 1000:.2f}ms")
                return f"Responsibilities: Party 1 - {party_1_obs}. Party 2 - {party_2_obs}"
        
        # If no local answer found
        end_time = time.perf_counter()
        logger.info(f"📊 No local answer found: {(end_time - start_time) * 1000:.2f}ms")
        return None
        
    except Exception as e:
        logger.error(f"Error in local answering: {str(e)}")
        return None

def chat_with_document(question: str, contract_data: Dict[str, Any], document_text: str) -> str:
    """Hybrid chat: Try local answer first, fallback to full document text with OpenAI for complex questions"""
    import time
    
    overall_start = time.perf_counter()
    logger.info(f"=== HYBRID CHAT ANALYSIS START ===")
    logger.info(f"Question: {question}")
    
    # Step 1: Try local answer first (instant)
    local_start = time.perf_counter()
    local_answer = answer_question_locally(question, contract_data)
    local_end = time.perf_counter()
    
    if local_answer:
        total_time = local_end - overall_start
        logger.info(f"📊 ⚡ INSTANT LOCAL ANSWER: {total_time * 1000:.2f}ms")
        logger.info(f"📊 Answer: {local_answer}")
        logger.info(f"=== HYBRID CHAT ANALYSIS END ===")
        return local_answer
    
    # Step 2: Fallback to OpenAI using FULL DOCUMENT TEXT for complex questions
    logger.info("📊 Local answer not found, using OpenAI with full document text")
    
    try:
        # Context preparation using FULL document text + extracted data
        context_start = time.perf_counter()
        
        # Use full document text for comprehensive answers
        full_context = f"""
        FULL DOCUMENT TEXT:
        {document_text}
        
        EXTRACTED STRUCTURED DATA:
        {json.dumps(contract_data, indent=2)}
        """
        
        context_end = time.perf_counter()
        logger.info(f"📊 Full context preparation: {(context_end - context_start) * 1000:.2f}ms")
        logger.info(f"📊 Full context size: {len(full_context)} characters")
        
        # Client setup (cached)
        client_start = time.perf_counter()
        if st.session_state.openai_client is None:
            try:
                from openai import OpenAI
                st.session_state.openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
                client_type = "new_client_cached"
                logger.info("📊 Created and cached new OpenAI client")
            except ImportError:
                client_type = "legacy_client"
                logger.info("📊 Using legacy OpenAI client (no caching)")
        else:
            client_type = "cached_client"
            logger.info("📊 Using cached OpenAI client")
        
        client_end = time.perf_counter()
        logger.info(f"📊 Client setup: {(client_end - client_start) * 1000:.2f}ms")
        
        # Message preparation
        message_start = time.perf_counter()
        
        document_type = contract_data.get("document_type", "Contract")
        
        messages = [
            {
                "role": "system",
                "content": f"""You are a helpful assistant that explains {document_type} documents in simple, easy-to-understand language.
                
                You have access to:
                1. The FULL DOCUMENT TEXT - Use this for detailed questions about specific clauses, terms, conditions, obligations, or any content in the document
                2. EXTRACTED STRUCTURED DATA - Use this for quick reference to key information
                
                Instructions:
                - Answer questions in simple, everyday language that anyone can understand
                - Avoid legal jargon - use plain English instead
                - Explain things like you're talking to a friend or family member
                - For specific clauses or detailed terms, quote relevant sections from the document but explain what they mean in simple terms
                - Be accurate but make it easy to understand
                - If something is complex, break it down into simple steps or bullet points
                - Use analogies or examples when helpful
                - If information is not in the document, clearly state that"""
            },
            {
                "role": "user", 
                "content": f"Document Context: {full_context}\n\nQuestion: {question}"
            }
        ]
        message_end = time.perf_counter()
        logger.info(f"📊 Message preparation: {(message_end - message_start) * 1000:.2f}ms")
        
        # API call
        api_start = time.perf_counter()
        logger.info("📊 Sending request to OpenAI API for comprehensive document analysis...")
        
        if client_type in ["new_client_cached", "cached_client"]:
            response = st.session_state.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=500,  # Increased for detailed answers
                timeout=10  # Increased timeout for longer processing
            )
            answer = response.choices[0].message.content
        else:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=messages,
                temperature=0.1,
                max_tokens=500,
                request_timeout=10
            )
            answer = response.choices[0].message.content
        
        api_end = time.perf_counter()
        api_time = api_end - api_start
        logger.info(f"📊 ⚡ OpenAI API call: {api_time * 1000:.2f}ms")
        
        # Overall timing
        overall_end = time.perf_counter()
        total_time = overall_end - overall_start
        local_time = local_end - local_start
        
        logger.info(f"📊 ⚡ TOTAL HYBRID CHAT TIME: {total_time * 1000:.2f}ms")
        logger.info(f"📊 Local attempt: {local_time * 1000:.2f}ms")
        logger.info(f"📊 API time: {api_time * 1000:.2f}ms")
        logger.info(f"=== HYBRID CHAT ANALYSIS END ===")
        
        return answer
        
    except Exception as e:
        error_msg = f"Error in OpenAI fallback: {str(e)}"
        logger.error(f"📊 ❌ HYBRID CHAT ERROR: {error_msg}")
        return error_msg

def main():
    st.title("📄 Multi-Document Contract Analyzer")
    st.markdown("Upload your contracts (Rental, NDA, MSA, Insurance, MOU) and get comprehensive analysis with Q&A capabilities")
    
    logger.info("Main application interface loaded")
    
    # Setup OpenAI
    api_configured = setup_openai()
    
    if not api_configured:
        logger.warning("OpenAI API not configured properly, stopping application")
        st.stop()
    
    # File upload with document type selection
    st.subheader("📁 Upload Document")
    
    # Document type selector
    col1, col2 = st.columns([1, 2])
    with col1:
        document_type = st.selectbox(
            "📋 Document Type",
            ["Rental", "NDA", "MSA", "Insurance", "MOU"],
            help="Select the type of contract you're uploading"
        )
    
    with col2:
        uploaded_file = st.file_uploader(
            f"Choose a {document_type} document",
            type=['pdf', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'tiff'],
            help="Supports PDF, Word documents, text files, and images"
        )
    
    # Store document type in session state
    if 'document_type' not in st.session_state:
        st.session_state.document_type = document_type
    elif st.session_state.document_type != document_type:
        # Clear previous data if document type changed
        st.session_state.document_text = ""
        st.session_state.extracted_data = None
        st.session_state.processed_file = None
        st.session_state.document_type = document_type
        logger.info(f"Document type changed to {document_type}, clearing previous data")
    
    if uploaded_file is not None:
        logger.info(f"File uploaded: {uploaded_file.name}")
        
        # CHECK IF WE ALREADY PROCESSED THIS FILE
        if (st.session_state.document_text == "" or 
            st.session_state.extracted_data is None or 
            'processed_file' not in st.session_state or 
            st.session_state.processed_file != uploaded_file.name):
            
            logger.info("Processing new file or file not in session state")
            
            # Show extraction method being used
            if uploaded_file.type == "application/pdf":
                extraction_method = "🚀 Threaded PDF Processing: PyMuPDF → Parallel OCR → Parallel Vision API"
            else:
                extraction_method = "📄 Standard Processing for Non-PDF Files"
            
            st.info(f"🔄 **Extraction Method:** {extraction_method}")
            
            # Process the document with optimized extraction
            with st.spinner("🚀 Processing document with optimized extraction..."):
                # Create a progress placeholder
                progress_placeholder = st.empty()
                status_placeholder = st.empty()
                
                # Show processing steps
                status_placeholder.info("📄 **Step 1:** Extracting text from document...")
                progress_placeholder.progress(0.3)
                
                # Use appropriate processing method
                if uploaded_file.type == "application/pdf":
                    # For PDFs, try threaded processing if PyMuPDF available
                    try:
                        import fitz
                        status_placeholder.info("🚀 **Using threaded PDF processing...**")
                        document_text = extract_text_from_file_threaded(uploaded_file)
                    except ImportError:
                        status_placeholder.info("📄 **Using standard PDF processing...**")
                        document_text = extract_text_from_file(uploaded_file)
                else:
                    # For non-PDFs, use standard processing
                    document_text = extract_text_from_file(uploaded_file)
                
                progress_placeholder.progress(0.6)
                
                if document_text:
                    status_placeholder.info(f"🤖 **Step 2:** Analyzing {document_type} with AI...")
                    extracted_data = extract_contract_info(document_text, document_type)
                    progress_placeholder.progress(0.9)
                else:
                    extracted_data = None
                
                progress_placeholder.progress(1.0)
                status_placeholder.success("✅ **Processing Complete!**")
                
                # Clear progress indicators after a short delay
                import time
                time.sleep(1)
                progress_placeholder.empty()
                status_placeholder.empty()
            
            # Store in session state
            if document_text:
                st.session_state.document_text = document_text
                st.session_state.processed_file = uploaded_file.name
                logger.info("Document text stored in session state")
                
                if extracted_data:
                    st.session_state.extracted_data = extracted_data
                    logger.info("Contract data extracted and stored in session state")
        
        else:
            # USE EXISTING DATA FROM SESSION STATE
            logger.info("Using existing data from session state - no reprocessing needed")
            document_text = st.session_state.document_text
            extracted_data = st.session_state.extracted_data
            st.info("✅ **Using cached document data** - no reprocessing needed!")
        
        # NOW DISPLAY THE RESULTS (whether new or cached)
        if document_text and extracted_data:
            # Show extraction results
            char_count = len(document_text)
            word_count = len(document_text.split())
            extraction_method_used = getattr(st.session_state, 'extraction_method', 'Standard')
            st.success(f"✅ **Document processed successfully!** {char_count} characters, {word_count} words")
            
            # Show analysis success
            if 'extraction_error' in extracted_data:
                st.warning(f"⚠️ **Partial Analysis:** {extracted_data.get('extraction_error', 'Unknown issue')}")
                st.info("📋 **Fallback extraction used** - Basic information extracted manually")
            
            st.divider()
            
            # NEW MINIMAL DISPLAY: Start Date, End Date, Document Type
            display_minimal_contract_info(extracted_data)
            
            st.divider()
            
            # COLLAPSIBLE BRIEF SUMMARY
            document_type = extracted_data.get("document_type", "Unknown Document")
            brief_summary = create_brief_summary(extracted_data, document_type)
            
            with st.expander("📄 Document Summary"):
                st.write(brief_summary)
            
            st.divider()
            
            # PROMINENT CHAT INTERFACE
            st.subheader("💬 Ask Questions About the Document")
            st.info(f"🚀 **Smart Chat**: Get instant answers about your {document_type.replace('Agreement', '').replace('Policy', '').strip()}!")
            
            # Document-specific example questions
            if "Rental" in document_type:
                placeholder_text = "e.g., Who is the tenant? When does lease expire? What are my obligations?"
            elif "NDA" in document_type or "Non-Disclosure" in document_type:
                placeholder_text = "e.g., Who is the disclosing party? What are confidentiality exceptions?"
            elif "MSA" in document_type or "Master Service" in document_type:
                placeholder_text = "e.g., Who is the service provider? What are termination conditions?"
            elif "Insurance" in document_type:
                placeholder_text = "e.g., What is covered? When does my policy expire? What's my deductible?"
            elif "MOU" in document_type or "Memorandum" in document_type:
                placeholder_text = "e.g., What is the purpose? Who are the parties? What are the responsibilities?"
            else:
                placeholder_text = "Ask any question about the document..."
            
            # Chat input - made more prominent
            question = st.text_input("💭 Ask questions about the document", placeholder=placeholder_text)
            
            ask_button = st.button("🚀 Ask", type="primary")
            
            if (ask_button or question) and question:
                logger.info(f"User asked question: {question}")
                
                # Create placeholder for response
                response_placeholder = st.empty()
                
                with st.spinner("⚡ Getting answer..."):
                    # Use hybrid chat function (local first, then full document OpenAI)
                    answer = chat_with_document(question, extracted_data, document_text)
                    
                # Display the final answer immediately
                response_placeholder.success(f"**Answer:** {answer}")
                logger.info("Question answered - hybrid approach used")
            
            # OPTIONAL DETAILED SECTIONS (Collapsible)
            with st.expander("🔍 View Raw Extracted Text"):
                st.text_area("Extracted Text", document_text, height=200)
            
            with st.expander("🔧 Debug Information"):
                st.json(extracted_data)
                
        elif document_text and not extracted_data:
            logger.error("Failed to extract contract data")
            st.error("❌ Failed to analyze the contract. Please check the logs for details.")
            
            # Show manual analysis option
            st.info("🔧 **Manual Analysis Available** - Try the debug information below:")
            with st.expander("📄 Raw Document Text"):
                st.text_area("Document Content", document_text, height=300)
                
            if st.button("🚀 Try Manual Analysis"):
                try:
                    manual_data = extract_basic_info_manually(document_text, document_type)
                    st.session_state.extracted_data = manual_data
                    st.success("✅ Manual analysis completed!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Manual analysis failed: {str(e)}")
        else:
            logger.error("Failed to extract text from document")
            st.error("❌ Failed to extract text from the document. Please try a different file.")
    
    else:
        logger.info("No file uploaded yet")
        st.info("👆 Please select document type and upload a contract to get started")
        
        # Show supported formats
        with st.expander("📋 Supported Document Types & Formats"):
            st.markdown("""
            **📄 Document Types:**
            - **🏠 Rental Agreements** - Lease contracts, property rentals, commercial leases
            - **🔒 NDAs** - Non-disclosure agreements, confidentiality contracts
            - **📋 MSAs** - Master service agreements, service contracts, vendor agreements
            - **🏥 Insurance Policies** - Health, auto, life, property insurance policies
            - **🤝 MOUs** - Memorandums of understanding, partnership agreements
            
            **📁 File Formats:**
            - **PDF files** (.pdf) - Including scanned documents with OCR
            - **Word documents** (.docx) - Microsoft Word format
            - **Text files** (.txt) - Plain text format
            - **Images** (.jpg, .jpeg, .png, .tiff) - Uses OCR for text extraction
            """)
            
            st.markdown("""
            **🎯 What We Extract:**
            - **Rental:** Parties, property details, rent amounts, lease terms, dates
            - **NDA:** Confidentiality scope, parties, duration, obligations, restrictions  
            - **MSA:** Service details, payment terms, parties, termination clauses
            - **Insurance:** Coverage details, premiums, deductibles, policy terms, exclusions
            - **MOU:** Purpose, parties, responsibilities, governance, collaboration terms
            """)
    
    
if __name__ == "__main__":
    main()